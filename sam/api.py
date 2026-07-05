from __future__ import annotations

import json
import os
import uuid
from datetime import datetime, timezone
from typing import Any, Literal, Optional
import base64
import json as jsonlib
import time
import hmac
import hashlib
import threading

from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import ec, rsa, padding
from cryptography.hazmat.primitives.asymmetric.utils import encode_dss_signature

from starlette.responses import JSONResponse, StreamingResponse

from fastapi import FastAPI, HTTPException, Response, Header, Request, Depends
from fastapi import UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.exc import OperationalError
import requests
import io
import csv
import re
import queue
import unicodedata
from difflib import SequenceMatcher
from pypdf import PdfReader

from .cache import (
    cache_clear_classify,
    cache_classify_is_disabled,
    cache_classify_set_disabled,
    cache_get,
    cache_set,
)
from .db import get_db
from .brand_messaging import INDICATIVE_DISCLAIMER_FR
from .classification_completeness import apply_completeness_adjustments, sanitize_provisional_narrative
from .classification_risk import enrich_classifications_with_risk
from .rgi import apply_rgi_pipeline_to_response
from .description_quality import assess_description_quality, enrich_item_description_quality
from .tariff_labels import (
    build_heading_narrative_index,
    build_tariff_label_index,
    lookup_position_label,
    resolve_hs_code_to_tec,
    set_heading_narrative_index,
    set_tariff_label_index,
)
from .tariff_metadata import get_full_chapter_name, get_full_section_name, get_position_heading
from .tariff_rates import build_tariff_rate_index, enrich_item_tariff_rates, set_tariff_rate_index
from .tariff_notes import (
    build_chapter_notes_index,
    build_chapter_titles_index,
    set_chapter_notes_index,
    set_chapter_titles_index,
)
from .tariff_position_rules import build_surface_sensitive_positions, set_surface_sensitive_positions
from .rag import (
    ClassificationPipelineResult,
    add_validated_classification_example_to_index,
    build_assistant_meta_response_json,
    initialize_chatbot,
    initialize_validated_classifications_index,
    is_assistant_meta_query,
    is_ui_boilerplate_line,
    process_user_input,
)
from .config.settings import Config
from .product_identification import product_identification_enabled
from .openai_web_search import openai_web_search_enabled
from .classification_progress import ClassificationProgressReporter, sse_event, sse_init_event
from .app_logger import get_logger

logger = get_logger(__name__)

_ALIAS_CACHE_LOCK = threading.Lock()
_ALIAS_CACHE: dict[str, Any] = {"fetched_at": 0.0, "aliases": {}}
_ALIAS_CACHE_TTL_SECONDS = 60.0
_CLASSIFY_CACHE_SCHEMA_VERSION = "v2"

_ALIAS_FUZZY_THRESHOLD = 0.86  # seuil de similarite chaine->chaine (flou)


def _strip_accents_ascii(s: str) -> str:
    """
    Normalise en ASCII pour eviter les ecarts "téléphone" vs "telephone".
    """
    return unicodedata.normalize("NFKD", (s or "")).encode("ascii", "ignore").decode("ascii")


def _default_aliases_map() -> dict[str, str]:
    return {
        # Uniquement des abréviations (pas de fautes en dur).
        "gsm": "telephone",
        "tel": "telephone",
        "pc": "ordinateur",
        "ordi": "ordinateur",
        # Canons identite: bootstrap pour le matching flou.
        # (On n'y met pas de fautes orthographiques spécifiques.)
        "canette": "canette",
        "bouteille": "bouteille",
        "cheval": "cheval",
        "bijou": "bijou",
    }


def _ensure_normalization_aliases_table() -> None:
    try:
        with get_db() as db:
            db.execute(
                text(
                    """
                    create table if not exists public.normalization_aliases (
                      id uuid primary key default gen_random_uuid(),
                      alias text not null unique,
                      canonical text not null,
                      is_active boolean not null default true,
                      created_at timestamptz not null default now(),
                      updated_at timestamptz not null default now()
                    )
                    """
                )
            )
            db.commit()
    except Exception:
        logger.warning("[alias] table normalization_aliases indisponible", exc_info=True)


def _load_aliases_map(refresh: bool = False) -> dict[str, str]:
    now = time.time()
    with _ALIAS_CACHE_LOCK:
        if not refresh:
            fetched_at = float(_ALIAS_CACHE.get("fetched_at") or 0.0)
            if _ALIAS_CACHE.get("aliases") and (now - fetched_at) < _ALIAS_CACHE_TTL_SECONDS:
                return dict(_ALIAS_CACHE["aliases"])

    # Toutes les clés/canoniques sont normalisées en minuscules ASCII
    # pour matcher avec `_normalize_item_key()` qui fait déjà `encode('ascii','ignore')`.
    aliases: dict[str, str] = {}
    for a, c in _default_aliases_map().items():
        a_n = _strip_accents_ascii(a).strip().lower()
        c_n = _strip_accents_ascii(c).strip().lower()
        if a_n and c_n:
            aliases[a_n] = c_n
    try:
        _ensure_normalization_aliases_table()
        with get_db() as db:
            rows = db.execute(
                text(
                    """
                    select alias, canonical
                    from public.normalization_aliases
                    where is_active = true
                    """
                )
            ).mappings().all()
            for row in rows:
                a = _strip_accents_ascii(str(row.get("alias") or "")).strip().lower()
                c = _strip_accents_ascii(str(row.get("canonical") or "")).strip().lower()
                if a and c:
                    aliases[a] = c
    except Exception:
        logger.warning("[alias] impossible de charger les alias BDD", exc_info=True)

    with _ALIAS_CACHE_LOCK:
        # Ajoute aussi un mapping identité pour chaque canon.
        # Utile pour que la partie "fuzzy" puisse reconnaître directement "ordinateur" etc.,
        # même si la canonnique n'est pas une clé d'alias en BDD.
        for canonical in list(aliases.values()):
            if canonical and canonical not in aliases:
                aliases[canonical] = canonical
        _ALIAS_CACHE["aliases"] = aliases
        _ALIAS_CACHE["fetched_at"] = now
    return dict(aliases)

def _classify_cache_key(query: str) -> str:
    """
    Génère une clé Redis courte et stable pour la classification.
    Evite les clés très longues (utile quand on classe depuis un fichier).
    """
    q = (query or "").strip().lower()
    digest = hashlib.sha256(q.encode("utf-8")).hexdigest()
    # Versionne la clé pour invalider les anciens résultats quand la logique
    # d'agrégation/normalisation évolue.
    return f"classify:{_CLASSIFY_CACHE_SCHEMA_VERSION}:{digest}"

# Dependency FastAPI pour endpoints admin.
# (La fonction `_require_admin` est définie plus bas ; ici on ne fait que
# déléguer, le nom est résolu au moment de l'appel.)
def admin_required(authorization: str | None = Header(default=None)) -> str:
    return _require_admin(authorization)


app = FastAPI(
    title="Mosam CEDEAO Classification API",
    version="0.1.0",
    description="API de classification tarifaire CEDEAO (RAG + OpenAI)",
)

_cors_extra = [
    o.strip()
    for o in (os.getenv("CORS_ALLOWED_ORIGINS") or "").split(",")
    if o.strip()
]
_cors_origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    *_cors_extra,
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    # Ports dev (3000, 3001, etc.) sans tout ouvrir en prod si tu complètes CORS_ALLOWED_ORIGINS.
    allow_origin_regex=r"https?://(localhost|127\.0\.0\.1)(:\d+)?$",
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _operational_error_client_detail(exc: OperationalError) -> str:
    """Message API selon la cause (pooler, IPv6, etc.)."""
    raw = str(getattr(exc, "orig", None) or exc)
    low = raw.lower()
    if "tenant or user not found" in low:
        return (
            "Supabase pooler: 'Tenant or user not found'. "
            "Copy the Session pooler URI from Dashboard → Connect (user postgres.YOUR_PROJECT_REF, "
            "host exactly as shown, port 5432). Do not use db.*.supabase.co for session mode."
        )
    if "timeout" in low or "10060" in raw or "timed out" in low:
        return (
            "Postgres connection timed out. Direct host db.*.supabase.co and port 6543 often resolve "
            "to IPv6 only. On IPv4-only networks use Session pooler: SUPABASE_DB_POOLER_URL with "
            "host aws-*.region.pooler.supabase.com and port 5432 (from Connect → Session pooler)."
        )
    return (
        "Database unavailable (Postgres connection failed). "
        "Prefer Session pooler in SUPABASE_DB_POOLER_URL (aws-*.pooler.supabase.com:5432, user "
        "postgres.projectref). Transaction pooler (db.*:6543) can still be IPv6-only."
    )


@app.exception_handler(OperationalError)
async def _handle_sqlalchemy_operational(_request: Request, exc: OperationalError) -> JSONResponse:
    """Connexion DB (timeout Supabase IPv6, etc.) : réponse JSON + CORS via middleware."""
    logger.error("Erreur SQLAlchemy OperationalError (souvent Postgres/Supabase)", exc_info=exc)
    return JSONResponse(
        status_code=503,
        content={"detail": _operational_error_client_detail(exc)},
    )


class MerchandiseItem(BaseModel):
    """Un article structuré tel que saisi dans le formulaire."""

    designation: str
    material: str = ""
    usage: str = ""
    characteristics: str = ""
    quantity: str = ""
    unit: str = ""
    origin: str = ""
    value: str = ""
    currency: str = ""


class ClassifyRequest(BaseModel):
    """Requête de classification d'une ou plusieurs marchandises."""

    query: str = ""
    items: list[MerchandiseItem] | None = None
    user_id: str | None = None


class ClassifyResponse(BaseModel):
    """Réponse brute renvoyée par le LLM (JSON sérialisé en texte)."""

    raw: str


class ClassifyFileResponse(BaseModel):
    """Réponse pour classification depuis un fichier."""

    raw: str
    # Texte/requête effectivement envoyée au moteur de classification.
    # Sert notamment au cache lors de la validation côté admin.
    effective_query: str
    items_count: int


class ValidateClassificationRequest(BaseModel):
    """
    Données envoyées par le frontend lorsqu'un agent
    valide une classification précise.
    """

    description: str
    section: str
    chapter: str
    hs_code: str
    confidence: float | None = None
    dd_rate: str | None = None
    rs_rate: str | None = None
    other_taxes: str | None = None
    us_unit: str | None = None
    origin: str | None = None
    value: str | None = None
    user_id: str | None = None
    quantity: int | None = None
    # Optionnel : si fournis, la réponse complète est mise en cache pour cette requête
    # (cache utilisé uniquement lorsqu'au moins une classification est validée)
    query: str | None = None
    raw_response: str | None = None
    # Optionnel : associe la classification à un "dossier" (ex: entreprise / société).
    # Si aucun dossier n'est fourni => classification non associée.
    dossier_name: str | None = None
    # Metadonnees Mosam (persistees a la validation pour l'historique et l'audit).
    justification: str | None = None
    risk_level: Literal["low", "medium", "high"] | None = None
    risk_label: str | None = None
    position_label: str | None = None
    classification_status: Literal["confirmee", "provisoire"] | None = None
    identification_confidence: float | None = None
    product_identification: dict[str, Any] | None = None
    source_query: str | None = None


_CLASSIFICATION_ENRICHMENT_SELECT = """
          c.justification,
          c.risk_level,
          c.risk_label,
          c.position_label,
          c.classification_mode,
          c.identification_confidence,
          c.product_identification,
          c.source_query,"""

_HISTORY_CSV_EXTRA_HEADERS = [
    "justification",
    "risk_level",
    "risk_label",
    "position_label",
    "classification_mode",
    "identification_confidence",
    "source_query",
]


def _classification_history_select_sql() -> str:
    return f"""
        select
          c.id,
          c.description_produit,
          c.section_produit,
          c.chapitre_produit,
          c.code_tarifaire,
          c.classification_confidence,
          c.quantity,
          c.dd_rate,
          c.rs_rate,
          c.other_taxes,
          c.us_unit,
          c.origin,
          c.value,
          c.statut_validation,
          c.created_at as date_classification,
          c.user_id,
{_CLASSIFICATION_ENRICHMENT_SELECT}
          d.id::text as dossier_id,
          d.name as dossier_name,
          u.nom_user as agent_name,
          u.id as agent_id
        from public.classifications c
        left join public.users u on u.id = c.user_id
        left join public.classification_dossiers d on d.id = c.dossier_id
    """


def _classification_enrichment_params(payload: ValidateClassificationRequest) -> dict[str, Any]:
    product_id = payload.product_identification
    product_id_json: str | None = None
    if isinstance(product_id, dict) and product_id:
        try:
            product_id_json = json.dumps(product_id, ensure_ascii=False)
        except Exception:
            product_id_json = None
    return {
        "justification": (payload.justification or "").strip() or None,
        "risk_level": payload.risk_level,
        "risk_label": (payload.risk_label or "").strip() or None,
        "position_label": (payload.position_label or "").strip() or None,
        "classification_mode": payload.classification_status,
        "identification_confidence": payload.identification_confidence,
        "product_identification": product_id_json,
        "source_query": (payload.source_query or "").strip() or None,
    }


class ValidateClassificationBulkRequest(BaseModel):
    """
    Validation de plusieurs classifications en une seule requête.

    Objectif : réduire drastiquement le nombre de calls HTTP (sinon le
    rate-limit sur `POST /classifications/validate` se déclenche).
    """

    items: list[ValidateClassificationRequest]
    # Optionnel : permet d'écrire la cache une seule fois pour toute la requête
    # même si l'UI ne fournit pas `query/raw_response` par item.
    query: str | None = None
    raw_response: str | None = None
    # Dossier appliqué à toutes les classifications validées dans cette requête.
    dossier_name: str | None = None


# Mapping chapitre SH (2 chiffres) -> (section romain, libellé section) pour corriger les réponses LLM
HS_CHAPTER_TO_SECTION: dict[int, tuple[str, str]] = {
    **{c: ("I", "Animaux vivants et produits du règne animal") for c in range(1, 6)},
    **{c: ("II", "Produits du règne végétal") for c in range(6, 15)},
    15: ("III", "Matières grasses et huiles ; cires"),
    **{c: ("IV", "Produits des industries alimentaires ; boissons ; tabacs") for c in range(16, 25)},
    **{c: ("V", "Produits minéraux") for c in range(25, 28)},
    **{c: ("VI", "Produits des industries chimiques") for c in range(28, 39)},
    **{c: ("VII", "Matières plastiques et caoutchouc") for c in range(39, 41)},
    **{c: ("VIII", "Peaux, cuirs, fourrures") for c in range(41, 44)},
    **{c: ("IX", "Bois, liège, vannerie") for c in range(44, 47)},
    **{c: ("X", "Pâtes de bois ; papier") for c in range(47, 50)},
    **{c: ("XI", "Textiles") for c in range(50, 64)},
    **{c: ("XII", "Chaussures ; coiffure ; parapluies") for c in range(64, 68)},
    **{c: ("XIII", "Pierre, plâtre, ciment ; verre ; céramique") for c in range(68, 72)},
    71: ("XIV", "Perles, pierres précieuses ; métaux précieux"),
    **{c: ("XV", "Métaux communs et ouvrages") for c in range(72, 84)},
    **{c: ("XVI", "Machines et appareils ; matériel électrique") for c in range(84, 86)},
    **{c: ("XVII", "Matériel de transport") for c in range(86, 90)},
    **{c: ("XVIII", "Instruments optiques, photographiques ; horlogerie ; instruments de précision") for c in range(90, 93)},
    93: ("XIX", "Armes et munitions"),
    **{c: ("XX", "Articles manufacturés divers") for c in range(94, 97)},
    **{c: ("XXI", "Œuvres d'art ; antiquités") for c in range(97, 100)},
}


def _normalize_section_chapter_from_hs(hs_code: str | None) -> dict[str, str]:
    """À partir d'un code SH (ex: 8517.13.00.00 ou 2008.19), déduit section et chapitre corrects."""
    if not hs_code or not isinstance(hs_code, str):
        return {}
    part = hs_code.strip().split(".")[0]
    if len(part) < 2:
        return {}
    try:
        ch = int(part[:2])
    except ValueError:
        return {}
    if ch < 1 or ch > 99:
        return {}
    section_roman, section_name = HS_CHAPTER_TO_SECTION.get(ch, ("", ""))
    full_section = get_full_section_name(section_roman, section_name)
    full_chapter = get_full_chapter_name(ch)
    return {
        "section": section_roman,
        "section_name": full_section,
        "chapter": f"{ch:02d}",
        "chapter_name": full_chapter,
    }


def _normalize_classifications_response(
    raw_text: str,
    progress: ClassificationProgressReporter | None = None,
    product_identifications: list[dict[str, Any]] | None = None,
) -> str:
    """
    Parse la réponse JSON du LLM, corrige section/chapitre à partir du code SH pour chaque
    classification, puis resérialise. Limite le nombre de lignes pour éviter les décompositions abusives.
    """
    if not raw_text or not raw_text.strip():
        return raw_text
    stripped = raw_text.strip()
    if stripped.startswith("```"):
        stripped = stripped.strip("`")
        if stripped.lower().startswith("json"):
            stripped = stripped[4:]
        stripped = stripped.strip("` \n\r")
    try:
        data = json.loads(stripped)
    except Exception:
        return raw_text
    if not isinstance(data, dict):
        return raw_text
    classifications = data.get("classifications")
    if not isinstance(classifications, list):
        return raw_text

    from .candidate_set_enforcer import attach_candidates_to_classifications

    attach_candidates_to_classifications(classifications, product_identifications)

    for item in classifications:
        if not isinstance(item, dict):
            continue
        hs = item.get("hs_code")
        if not hs:
            continue
        source_text = str(item.get("source_query") or item.get("description") or "")
        resolved_hs = resolve_hs_code_to_tec(str(hs), description=source_text)
        if resolved_hs and resolved_hs != str(hs).strip():
            item["hs_code"] = resolved_hs
            hs = resolved_hs
        s = str(hs).strip().upper()
        if s in ("NON APPLICABLE", "NON RENSEIGNÉ", "N/A", "NA") or len(s) < 4:
            continue
        if not s.replace(".", "").replace(" ", "").isdigit():
            continue
        normalized = _normalize_section_chapter_from_hs(str(hs))
        if normalized:
            # Corrige systematiquement section/chapitre a partir du code SH (le LLM se trompe souvent).
            if normalized.get("section"):
                item["section"] = normalized["section"]
            if normalized.get("section_name"):
                item["section_name"] = normalized["section_name"]
            if normalized.get("chapter"):
                item["chapter"] = normalized["chapter"]
        position_label = lookup_position_label(str(hs))
        if position_label and not item.get("position_label"):
            item["position_label"] = position_label
    # Nettoyage d'affichage : certaines zones peuvent contenir des caractères
    # non parfaitement supportés par la chaîne UI/copie-collage.
    # On supprime les accents pour garantir un rendu stable.
    text_keys = [
        "narrative",
        "description",
        "justification",
        "excerpt",
        "position_label",
        "section_name",
        "chapter_name",
        "subposition_label",
        "origin",
        "value",
        "dd_rate",
        "rs_rate",
        "us_unit",
        "other_taxes",
        "taxes_note",
    ]

    # Réponses `assistant_info` : garder les accents du narrative (texte fixe / présentation).
    if (
        isinstance(data.get("narrative"), str)
        and data.get("narrative")
        and not data.get("assistant_info")
    ):
        data["narrative"] = _strip_accents_ascii(data["narrative"])

    for item in classifications:
        if not isinstance(item, dict):
            continue
        for k in text_keys:
            if k == "narrative":
                continue
            if isinstance(item.get(k), str) and item.get(k):
                item[k] = _strip_accents_ascii(item[k])

    data["classifications"] = _filter_phantom_classifications(classifications)

    if not data["classifications"] and product_identifications:
        for pid in product_identifications:
            if not isinstance(pid, dict) or pid.get("skipped"):
                continue
            ptype = str(pid.get("product_type") or "").strip()
            pname = str(pid.get("product_name") or pid.get("original_query") or "").strip()
            if ptype or pname:
                data["classifications"].append({
                    "description": pname or ptype,
                    "hs_code": "",
                    "confidence": 20,
                    "classification_status": "provisoire",
                    "justification": (
                        f"Classification automatique non resolue. "
                        f"Produit identifie : {ptype or pname}. "
                        f"Preciser la description pour obtenir un code SH."
                    ),
                    "product_identification": pid,
                })
                break

    from .functional_coherence import apply_functional_coherence_gate

    for idx, item in enumerate(data["classifications"]):
        if not isinstance(item, dict):
            continue
        prod_id = None
        if product_identifications and idx < len(product_identifications):
            prod_id = product_identifications[idx]
        elif isinstance(item.get("product_identification"), dict):
            prod_id = item["product_identification"]
        apply_functional_coherence_gate(item, prod_id)

    from .position_validator import apply_position_validation

    for idx, item in enumerate(data["classifications"]):
        if not isinstance(item, dict):
            continue
        prod_id = None
        if product_identifications and idx < len(product_identifications):
            prod_id = product_identifications[idx]
        elif isinstance(item.get("product_identification"), dict):
            prod_id = item["product_identification"]
        candidates = item.get("tec_position_candidates")
        apply_position_validation(item, prod_id, candidates)

    if progress:
        progress.start("subposition")

    for idx, item in enumerate(data["classifications"]):
        if isinstance(item, dict):
            source = item.get("source_query") or item.get("description") or ""
            prod_id = None
            if product_identifications and idx < len(product_identifications):
                prod_id = product_identifications[idx]
            if isinstance(prod_id, dict) and not prod_id.get("skipped"):
                enriched = str(prod_id.get("enriched_description") or "").strip()
                ptype = str(prod_id.get("product_type") or "").strip()
                fusage = str(prod_id.get("function_usage") or "").strip()
                extra = " ".join(filter(None, [enriched, ptype, fusage]))
                if extra and extra.casefold() != source.casefold():
                    source = f"{source}\n{extra}".strip()
            from .classification_completeness import apply_early_subposition_gate

            apply_early_subposition_gate(item, source_text=source)
    if progress:
        progress.complete("subposition")
        progress.start("rgi")
    data = apply_rgi_pipeline_to_response(data)
    if progress:
        progress.complete("rgi")
    for item in data["classifications"]:
        if isinstance(item, dict):
            source = item.get("source_query") or item.get("description")
            enrich_item_description_quality(item, source_text=source)
            apply_completeness_adjustments(item, source_text=source)
            hs = item.get("hs_code")
            if hs:
                normalized = _normalize_section_chapter_from_hs(str(hs))
                if normalized:
                    if normalized.get("section"):
                        item["section"] = normalized["section"]
                    if normalized.get("section_name"):
                        item["section_name"] = normalized["section_name"]
                    if normalized.get("chapter"):
                        item["chapter"] = normalized["chapter"]
                    item["chapter_name"] = get_full_chapter_name(
                        normalized.get("chapter") or item.get("chapter") or "",
                        str(item.get("chapter_name") or ""),
                    )
                if item.get("subposition_status") == "a_determiner":
                    heading = get_position_heading(str(hs))
                    if heading:
                        item["position_label"] = heading
    if progress:
        progress.start("duties")
    for item in data["classifications"]:
        if isinstance(item, dict):
            enrich_item_tariff_rates(item)
    if progress:
        progress.complete("duties")
        progress.start("report")
    if isinstance(data.get("narrative"), str) and data.get("narrative"):
        data["narrative"] = sanitize_provisional_narrative(
            data["narrative"],
            [item for item in data["classifications"] if isinstance(item, dict)],
        )
    enrich_classifications_with_risk(data["classifications"])
    if progress:
        progress.complete("report")

    return json.dumps(data, ensure_ascii=False)


def _inject_source_query_into_llm_response(raw_text: str, source_query: str) -> str:
    """Attache le texte source utilisateur a chaque classification pour l'analyse de completude."""
    if not raw_text or not source_query:
        return raw_text
    stripped = raw_text.strip()
    if stripped.startswith("```"):
        stripped = stripped.strip("`")
        if stripped.lower().startswith("json"):
            stripped = stripped[4:]
        stripped = stripped.strip("` \n\r")
    try:
        data = json.loads(stripped)
    except Exception:
        return raw_text
    if not isinstance(data, dict):
        return raw_text
    classifications = data.get("classifications")
    if not isinstance(classifications, list):
        return raw_text
    for item in classifications:
        if isinstance(item, dict):
            item["source_query"] = source_query
    return json.dumps(data, ensure_ascii=False)


def _base_description_for_merge(description: str) -> str:
    """
    Canonise une description pour agréger des doublons LLM.
    Objectif: "X (variante ...)" et "X" doivent tomber sur la même clé.
    """
    s = (description or "").strip()
    # Supprime le contenu entre parenthèses (souvent les variantes orthographiques/infos additionnelles).
    s = re.sub(r"\s*\([^)]*\)\s*", " ", s).strip()
    s = re.sub(r"\s+", " ", s).strip()

    # Canonise davantage en réutilisant la normalisation de clé.
    # Objectif : "canettes" vs "cannettes", accents, casse, petits liens linguistiques.
    norm = _normalize_item_key(s)
    return norm or s.lower()


def _merge_duplicate_classifications(classifications: list[Any]) -> list[dict[str, Any]]:
    """
    Fusionne des classifications en doublon renvoyées par le LLM.
    Règle: même `hs_code` et même description "de base".
    """
    merged_by_key: dict[tuple[str, str], dict[str, Any]] = {}

    def _safe_int(v: Any, default: int) -> int:
        try:
            return int(v)
        except Exception:
            return default

    for cls in classifications:
        if not isinstance(cls, dict):
            continue
        hs_code = str(cls.get("hs_code") or "").strip()
        desc = str(cls.get("description") or "").strip()
        base_desc = _base_description_for_merge(desc)
        if not hs_code or not base_desc:
            continue

        key = (hs_code, base_desc)
        if key not in merged_by_key:
            # Stocke une copie superficielle pour éviter de muter l'entrée d'origine hors contrôle.
            merged_by_key[key] = dict(cls)
            continue

        target = merged_by_key[key]
        qty_existing = _safe_int(target.get("quantity"), 0)
        qty_new = _safe_int(cls.get("quantity"), 0)
        if qty_existing < 1:
            qty_existing = 0
        if qty_new < 1:
            qty_new = 0
        total_qty = qty_existing + qty_new

        target["quantity"] = total_qty if total_qty > 0 else target.get("quantity")

        # Harmonise la source: si plusieurs sources contribuent -> mixte.
        src_existing = str(target.get("quantity_source") or "").strip() or "explicit"
        src_new = str(cls.get("quantity_source") or "").strip() or "explicit"
        target["quantity_source"] = "mixte" if src_existing != src_new else src_existing

        # Concatène/dédouble les raw samples quand c'est présent.
        raw_existing = str(target.get("quantity_raw") or "")
        raw_new = str(cls.get("quantity_raw") or "")
        if raw_existing or raw_new:
            samples: list[str] = []
            for raw in [raw_existing, raw_new]:
                for part in raw.split(","):
                    p = part.strip()
                    if p and p not in samples:
                        samples.append(p)
            target["quantity_raw"] = ", ".join(samples[:10])

        # Confidence: moyenne pondérée par quantité.
        conf_existing = _safe_int(target.get("quantity_confidence"), 0)
        conf_new = _safe_int(cls.get("quantity_confidence"), 0)
        if total_qty > 0 and (conf_existing or conf_new):
            conf_sum = conf_existing * qty_existing + conf_new * qty_new
            target["quantity_confidence"] = int(round(conf_sum / total_qty))

        # Garde la première justification/sections (souvent similaires) : pas d'assemblage agressif.
        if (not target.get("excerpt")) and cls.get("excerpt"):
            target["excerpt"] = cls.get("excerpt")
        if (not target.get("justification")) and cls.get("justification"):
            target["justification"] = cls.get("justification")

    # Retourne dans l'ordre d'apparition des premières occurrences.
    order_keys = []
    for cls in classifications:
        if not isinstance(cls, dict):
            continue
        hs_code = str(cls.get("hs_code") or "").strip()
        base_desc = _base_description_for_merge(str(cls.get("description") or "").strip())
        key = (hs_code, base_desc)
        if key in merged_by_key and key not in order_keys:
            order_keys.append(key)

    return [merged_by_key[k] for k in order_keys]


def _extract_classifications(raw_text: str) -> list[dict]:
    """
    Essaie d'extraire la liste `classifications` à partir d'une chaîne brute.

    Gère plusieurs cas :
    - JSON direct: {"narrative": "...", "classifications": [...]}
    - JSON encodé en string: "\"{...}\""
    - Présence éventuelle de fences ```json ... ``` autour du JSON.
    """

    def _strip_fences(s: str) -> str:
        t = s.strip()
        if t.startswith("```"):
            # enlève ```json ou ``` puis la clôture finale ```
            t = t.strip("`")
            # si ça commence par json après les backticks
            if t.lower().startswith("json"):
                t = t[4:]
            return t.strip("` \n\r")
        return t

    current: object = raw_text
    for _ in range(3):
        if isinstance(current, str):
            candidate = _strip_fences(current)
            try:
                current = json.loads(candidate)
            except Exception:
                return []

        if isinstance(current, dict):
            classifications = current.get("classifications") or []
            return classifications if isinstance(classifications, list) else []
    return []


def _try_repair_json(s: str) -> dict | None:
    """Tente de réparer un JSON tronqué en fermant les délimiteurs manquants."""
    s = s.strip()
    if not s:
        return None
    if s.startswith("```"):
        s = s.strip("`")
        if s.lower().startswith("json"):
            s = s[4:]
        s = s.strip("` \n\r")

    open_braces = s.count("{") - s.count("}")
    open_brackets = s.count("[") - s.count("]")
    if open_braces <= 0 and open_brackets <= 0:
        return None
    repaired = s.rstrip(", \n\r\t")
    repaired += "]" * max(open_brackets, 0)
    repaired += "}" * max(open_braces, 0)
    try:
        obj = json.loads(repaired)
        if isinstance(obj, dict):
            return obj
    except Exception:
        pass
    return None


def _ensure_json_raw(raw: Any) -> str:
    """
    Force `raw` à être une chaîne JSON valide (contrat backend => frontend).

    IMPORTANT (contrat strict) :
    - On renvoie *toujours* une string JSON valide, même si le contenu original
      n'est pas un JSON parsable.
    """
    if isinstance(raw, (dict, list)):
        return json.dumps(raw, ensure_ascii=False)

    s = str(raw).strip()
    if not s:
        return json.dumps({"error": "empty_raw"}, ensure_ascii=False)

    # 1) Cas standard : raw est du JSON valide
    try:
        obj: object = json.loads(s)
    except Exception:
        repaired = _try_repair_json(s)
        if repaired is not None:
            return json.dumps(repaired, ensure_ascii=False)
        return json.dumps(
            {"error": "invalid_json", "raw_preview": s[:200]},
            ensure_ascii=False,
        )

    # 3) Cas : `raw` est une string qui contient (encore) du JSON
    if isinstance(obj, str):
        inner = obj.strip()
        try:
            inner_obj = json.loads(inner)
            return json.dumps(inner_obj, ensure_ascii=False)
        except Exception:
            return json.dumps({"error": "json_string_not_json", "raw": inner[:500]}, ensure_ascii=False)

    # 4) Tout le reste : on encapsule l'objet JSON parsé
    try:
        return json.dumps(obj, ensure_ascii=False)
    except Exception:
        return json.dumps({"error": "json_serialize_failed", "raw_preview": s[:200]}, ensure_ascii=False)


def _unwrap_pipeline_result(
    result: ClassificationPipelineResult | str,
) -> ClassificationPipelineResult:
    if isinstance(result, ClassificationPipelineResult):
        return result
    return ClassificationPipelineResult(llm_raw=str(result))


def _attach_product_identification(
    data: dict[str, Any],
    product_identifications: list[dict[str, Any]] | None,
) -> None:
    if not product_identifications:
        return
    active = [
        entry
        for entry in product_identifications
        if isinstance(entry, dict) and not entry.get("skipped")
    ]
    if not active:
        return
    data["product_identification"] = active
    classifications = data.get("classifications")
    if not isinstance(classifications, list):
        return
    for index, item in enumerate(classifications):
        if not isinstance(item, dict) or index >= len(product_identifications):
            continue
        entry = product_identifications[index]
        if isinstance(entry, dict) and not entry.get("skipped"):
            item["product_identification"] = entry
            web_sources = entry.get("web_sources") or []
            if isinstance(web_sources, list) and web_sources:
                item["web_sources"] = web_sources
            web_search_queries = entry.get("web_search_queries") or []
            if isinstance(web_search_queries, list) and web_search_queries:
                item["web_search_queries"] = web_search_queries
            if entry.get("web_search_used"):
                item["web_search_used"] = True
            candidates = entry.get("tec_position_candidates")
            if isinstance(candidates, list) and candidates:
                item["tec_position_candidates"] = candidates
            enriched = str(entry.get("enriched_description") or "").strip()
            existing = str(item.get("source_query") or "").strip()
            if enriched and not (
                _is_structured_product_dossier_text(existing)
                or re.search(r"(?i)\borigine\s*:", existing)
                or re.search(r"(?i)\bvaleur\s*:", existing)
            ):
                item["source_query"] = enriched

            id_conf = int(entry.get("identification_confidence") or 100)
            cls_conf = int(item.get("confidence") or item.get("classification_confidence") or 95)
            final_conf = min(id_conf, cls_conf)
            item["identification_confidence"] = id_conf
            item["classification_confidence"] = cls_conf
            item["confidence"] = final_conf


def _finalize_classification_response(
    raw_text: str,
    product_identifications: list[dict[str, Any]] | None = None,
    progress: ClassificationProgressReporter | None = None,
) -> str:
    normalized = _normalize_classifications_response(
        raw_text,
        progress=progress,
        product_identifications=product_identifications,
    )
    if not product_identifications:
        return normalized
    try:
        data = json.loads(_ensure_json_raw(normalized))
    except Exception:
        return normalized
    if isinstance(data, dict):
        _attach_product_identification(data, product_identifications)
        classifications = data.get("classifications")
        if isinstance(classifications, list):
            enrich_classifications_with_risk(classifications)
        return _ensure_json_raw(data)
    return normalized


def _inspect_raw_json(raw_out: str, request_id: str, when: str) -> None:
    """
    Logs de diagnostic : raw_out est censé être une string JSON.
    Indique si JSON.parse marche + existence/taille de `classifications`.
    """
    try:
        parsed = json.loads(raw_out)
    except Exception as e:
        logger.debug(
            "[classify %s] inspect %s: JSON.parse failed: %s", request_id, when, e
        )
        return

    if not isinstance(parsed, dict):
        logger.debug(
            "[classify %s] inspect %s: parsed not dict: %s",
            request_id,
            when,
            type(parsed),
        )
        return

    narrative_ok = isinstance(parsed.get("narrative"), str) and bool(parsed.get("narrative"))
    classifications = parsed.get("classifications")
    cls_ok = isinstance(classifications, list)
    cls_len = len(classifications) if cls_ok else -1
    logger.debug(
        "[classify %s] inspect %s: narrative_ok=%s classifications_ok=%s classifications_len=%s",
        request_id,
        when,
        narrative_ok,
        cls_ok,
        cls_len,
    )


class UserCreate(BaseModel):
    """Données nécessaires à la création d'un utilisateur simple."""

    nom_user: str
    identifiant_user: str
    email: str
    is_admin: bool = False


class UserUpdate(BaseModel):
    """
    Données autorisées pour la mise à jour d'un utilisateur.

    Tous les champs sont facultatifs pour permettre des PATCH partiels.
    """

    nom_user: Optional[str] = None
    identifiant_user: Optional[str] = None
    email: Optional[str] = None
    is_admin: Optional[bool] = None
    statut: Optional[Literal["actif", "inactif", "supprimé"]] = None


class ResetPasswordResponse(BaseModel):
    """Réponse renvoyée après une réinitialisation de mot de passe."""

    user_id: str
    email: str
    new_password: str


class ClassificationStatusUpdate(BaseModel):
    """
    Données envoyées lorsqu'un admin force le statut d'une classification.
    """

    statut_validation: Literal["validé", "invalidé", "archivé"]


class ClassificationQuantityUpdate(BaseModel):
    """
    Données envoyées lorsqu'un admin corrige la quantité d'une classification.
    """

    quantity: int


class NormalizationAliasCreate(BaseModel):
    alias: str
    canonical: str
    is_active: bool = True


class NormalizationAliasUpdate(BaseModel):
    canonical: str | None = None
    is_active: bool | None = None


class AuditLogItem(BaseModel):
    id: str
    created_at: datetime
    actor_id: str
    action: str
    entity_type: str
    entity_id: str
    details: dict[str, Any] | None = None


@app.patch(
    "/classifications/{classification_id}/status",
    tags=["classification"],
)
def update_classification_status(
    classification_id: str,
    payload: ClassificationStatusUpdate,
    admin_id: str = Depends(admin_required),
) -> dict:
    """
    Met à jour le statut de validation d'une classification (admin uniquement).
    """

    actor_id = admin_id

    with get_db() as db:
        row = db.execute(
            text(
                """
                update public.classifications
                set statut_validation = :statut
                where id::text = :classification_id
                returning
                  id,
                  description_produit,
                  section_produit,
                  chapitre_produit,
                  code_tarifaire,
                  classification_confidence,
                  quantity,
                  dd_rate,
                  rs_rate,
                  other_taxes,
                  us_unit,
                  origin,
                  value,
                  user_id,
                  statut_validation,
                  created_at as date_classification
                """
            ),
            {
                "classification_id": classification_id,
                "statut": payload.statut_validation,
            },
        ).mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail="Classification introuvable.")

        db.commit()
        result = dict(row)

    _insert_audit_log(
        actor_id=actor_id,
        action="classification.update_status",
        entity_type="classification",
        entity_id=str(classification_id),
        details={
            "new_statut_validation": payload.statut_validation,
        },
    )

    return result


@app.patch(
    "/classifications/{classification_id}/quantity",
    tags=["classification"],
)
def update_classification_quantity(
    classification_id: str,
    payload: ClassificationQuantityUpdate,
    admin_id: str = Depends(admin_required),
) -> dict:
    """
    Met à jour la quantité d'une classification (admin uniquement).
    """
    if payload.quantity < 1:
        raise HTTPException(status_code=400, detail="La quantité doit être >= 1.")

    with get_db() as db:
        row = db.execute(
            text(
                """
                update public.classifications
                set quantity = :quantity
                where id::text = :classification_id
                returning
                  id,
                  description_produit,
                  section_produit,
                  chapitre_produit,
                  code_tarifaire,
                  classification_confidence,
                  quantity,
                  dd_rate,
                  rs_rate,
                  other_taxes,
                  us_unit,
                  origin,
                  value,
                  user_id,
                  statut_validation,
                  created_at as date_classification
                """
            ),
            {
                "classification_id": classification_id,
                "quantity": payload.quantity,
            },
        ).mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail="Classification introuvable.")

        db.commit()
        result = dict(row)

    _insert_audit_log(
        actor_id=admin_id,
        action="classification.update_quantity",
        entity_type="classification",
        entity_id=str(classification_id),
        details={"new_quantity": payload.quantity},
    )

    return result


def _decode_supabase_jwt(token: str) -> dict | None:
    """
    Décodage best-effort du payload JWT Supabase (sans vérification de signature).

    On utilise uniquement cette info pour identifier l'utilisateur courant côté
    backend lorsque la requête provient de notre frontend de confiance.
    """

    try:
        parts = token.split(".")
        if len(parts) != 3:
            return None
        payload_b64 = parts[1]
        # Ajoute le padding manquant si nécessaire
        padding = "=" * (-len(payload_b64) % 4)
        payload_bytes = base64.urlsafe_b64decode(payload_b64 + padding)
        return jsonlib.loads(payload_bytes.decode("utf-8"))
    except Exception:
        return None


_RATE_LIMIT_LOCK = threading.Lock()
_RATE_LIMIT_STATE: dict[str, list[float]] = {}


def _rate_limit(request: Request | None, scope: str, limit: int = 20, window_seconds: int = 60) -> None:
    """
    Limiteur simple par IP (Upstash Redis en prod, sinon mémoire locale).
    Empêche le flood sur endpoints sensibles.
    """
    ip = "unknown"
    try:
        if request and request.client and request.client.host:
            ip = request.client.host
    except Exception:
        pass

    key = f"ratelimit:{scope}:{ip}"

    # 1) Tentative : Upstash Redis (comportement stable multi-process)
    upstash_url = Config.UPSTASH_REDIS_REST_URL
    upstash_token = Config.UPSTASH_REDIS_REST_TOKEN
    if upstash_url and upstash_token:
        try:
            resp = requests.post(
                upstash_url,
                headers={"Authorization": f"Bearer {upstash_token}"},
                json=["INCR", key],
                timeout=5,
            )
            if resp.ok:
                data = resp.json()
                count = data.get("result")
                count_i: int | None = None
                if isinstance(count, (int, float)):
                    count_i = int(count)
                elif isinstance(count, str) and count.isdigit():
                    count_i = int(count)

                if count_i is not None:
                    if count_i == 1:
                        # TTL sur la première requête pour initier la fenêtre.
                        requests.post(
                            upstash_url,
                            headers={"Authorization": f"Bearer {upstash_token}"},
                            json=["EXPIRE", key, window_seconds],
                            timeout=5,
                        )
                    if count_i > limit:
                        raise HTTPException(
                            status_code=429,
                            detail="Trop de requêtes. Réessayez plus tard.",
                        )
                    return
        except HTTPException:
            raise
        except Exception:
            # Fallback : si Upstash est indisponible, on passe en mémoire locale.
            logger.warning("[rate_limit] Upstash indisponible, fallback mémoire locale")

    # 2) Fallback : mémoire locale (dev)
    now = time.time()
    with _RATE_LIMIT_LOCK:
        timestamps = _RATE_LIMIT_STATE.get(key, [])
        # Nettoyage fenêtre
        timestamps = [t for t in timestamps if now - t < window_seconds]
        if len(timestamps) >= limit:
            raise HTTPException(status_code=429, detail="Trop de requêtes. Réessayez plus tard.")
        timestamps.append(now)
        _RATE_LIMIT_STATE[key] = timestamps


def _base64url_decode(s: str) -> bytes:
    padding = "=" * (-len(s) % 4)
    return base64.urlsafe_b64decode(s + padding)


_JWKS_CACHE_LOCK = threading.Lock()
_JWKS_CACHE: dict[str, Any] = {"fetched_at": 0.0, "keys": []}


def _get_expected_supabase_iss() -> str:
    # Exemple d'issuer attendu (Supabase) :
    # https://<ref>.supabase.co/auth/v1
    return f"{Config.SUPABASE_URL.rstrip('/')}/auth/v1"


def _fetch_supabase_jwks() -> list[dict[str, Any]]:
    """
    Récupère la JWKS pour vérifier les JWT Supabase asymétriques (ES256/RS256).
    """
    jwks_url = f"{Config.SUPABASE_URL.rstrip('/')}/auth/v1/.well-known/jwks.json"
    resp = requests.get(jwks_url, timeout=5)
    if not resp.ok:
        raise RuntimeError(f"JWKS fetch failed: {resp.status_code}")
    data = resp.json()
    keys = data.get("keys")
    if not isinstance(keys, list):
        return []
    return keys


def _get_supabase_jwks_keys() -> list[dict[str, Any]]:
    ttl_seconds = 600.0
    now = time.time()
    with _JWKS_CACHE_LOCK:
        fetched_at = float(_JWKS_CACHE.get("fetched_at") or 0.0)
        if _JWKS_CACHE.get("keys") and (now - fetched_at) < ttl_seconds:
            return _JWKS_CACHE["keys"]
        try:
            keys = _fetch_supabase_jwks()
            _JWKS_CACHE["keys"] = keys
            _JWKS_CACHE["fetched_at"] = now
            return keys
        except Exception as exc:
            logger.warning("[auth] JWKS indisponible, vérification impossible: %s", exc)
            return []


def _b64url_to_int(s: str) -> int:
    return int.from_bytes(_base64url_decode(s), "big")


def _public_key_from_jwk(jwk: dict[str, Any]) -> rsa.RSAPublicKey | ec.EllipticCurvePublicKey | None:
    kty = jwk.get("kty")
    if kty == "EC":
        crv = str(jwk.get("crv") or "")
        if crv == "P-256":
            curve = ec.SECP256R1()
        elif crv == "P-384":
            curve = ec.SECP384R1()
        elif crv == "P-521":
            curve = ec.SECP521R1()
        else:
            return None

        x = _b64url_to_int(str(jwk.get("x") or ""))
        y = _b64url_to_int(str(jwk.get("y") or ""))
        numbers = ec.EllipticCurvePublicNumbers(x, y, curve)
        return numbers.public_key()

    if kty == "RSA":
        n = _b64url_to_int(str(jwk.get("n") or ""))
        e = _b64url_to_int(str(jwk.get("e") or ""))
        numbers = rsa.RSAPublicNumbers(e, n)
        return numbers.public_key()

    return None


def _verify_supabase_jwt(token: str) -> dict | None:
    """
    Vérifie la signature d'un JWT Supabase.

    - Si `alg` est symétrique (HS256/HS512) => vérifie avec `SUPABASE_JWT_SECRET` (HMAC).
    - Sinon (ES256/RS256...) => vérifie avec la JWKS Supabase (asymétrique).

    En cas de problème (secret absent, JWKS indisponible, alg inconnu, signature invalide),
    on refuse => `None`.
    """
    try:
        parts = token.split(".")
        if len(parts) != 3:
            return None

        header_b = _base64url_decode(parts[0])
        payload_b = _base64url_decode(parts[1])
        sig_b = _base64url_decode(parts[2])

        header = jsonlib.loads(header_b.decode("utf-8"))
        alg = str(header.get("alg") or "").upper()
        kid = header.get("kid")

        signing_input = f"{parts[0]}.{parts[1]}".encode("utf-8")

        secret = Config.SUPABASE_JWT_SECRET
        if alg in ("HS256", "HS512"):
            if not secret:
                logger.warning("[auth] HS secret manquant, refuser JWT (alg=%s)", alg)
                return None
            if alg == "HS256":
                expected = hmac.new(
                    secret.encode("utf-8"), signing_input, hashlib.sha256
                ).digest()
            else:
                expected = hmac.new(
                    secret.encode("utf-8"), signing_input, hashlib.sha512
                ).digest()

            if not hmac.compare_digest(expected, sig_b):
                logger.warning("[auth] JWT HS signature invalide (alg=%s)", alg)
                return None

        elif alg == "ES256":
            if not kid:
                logger.warning("[auth] JWT ES256 kid manquant")
                return None
            # ECDSA signature en JWS est `r||s` (64 bytes pour P-256)
            if len(sig_b) != 64:
                logger.warning(
                    "[auth] JWT ES256 signature invalide longueur=%s",
                    len(sig_b),
                )
                return None

            keys = _get_supabase_jwks_keys()
            jwk = next((k for k in keys if k.get("kid") == kid), None)
            if not jwk:
                logger.warning("[auth] JWKS clé ES256 introuvable (kid=%s)", kid)
                return None
            pub = _public_key_from_jwk(jwk)
            if not pub:
                logger.warning("[auth] JWKS public key ES256 invalide (kid=%s)", kid)
                return None

            r = int.from_bytes(sig_b[:32], "big")
            s = int.from_bytes(sig_b[32:], "big")
            signature_der = encode_dss_signature(r, s)
            pub.verify(signature_der, signing_input, ec.ECDSA(hashes.SHA256()))

        elif alg == "RS256":
            if not kid:
                logger.warning("[auth] JWT RS256 kid manquant")
                return None
            keys = _get_supabase_jwks_keys()
            jwk = next((k for k in keys if k.get("kid") == kid), None)
            if not jwk:
                logger.warning("[auth] JWKS clé RS256 introuvable (kid=%s)", kid)
                return None
            pub = _public_key_from_jwk(jwk)
            if not pub:
                logger.warning("[auth] JWKS public key RS256 invalide (kid=%s)", kid)
                return None

            pub.verify(
                sig_b,
                signing_input,
                padding.PKCS1v15(),
                hashes.SHA256(),
            )

        else:
            logger.warning("[auth] JWT alg non supporté: %s", alg)
            return None

        payload = jsonlib.loads(payload_b.decode("utf-8"))

        expected_iss = _get_expected_supabase_iss()
        iss = payload.get("iss")
        # Certains JWT peuvent avoir un trailing slash dans `iss`.
        # On normalise pour éviter les rejets "faussement invalides".
        if isinstance(iss, str) and iss.strip().rstrip("/") != expected_iss.strip().rstrip("/"):
            logger.warning("[auth] JWT issuer invalide (got=%s expected=%s)", iss, expected_iss)
            return None

        exp = payload.get("exp")
        if isinstance(exp, (int, float)) and exp < time.time():
            # Log prudent: ne pas divulguer le token, uniquement les timestamps.
            logger.warning("[auth] JWT expired exp=%s now=%s", exp, time.time())
            return None
        return payload
    except Exception:
        logger.warning("[auth] JWT vérification exception inattendue", exc_info=True)
        return None


def _require_admin(authorization: str | None) -> str:
    """
    Vérifie que le porteur du JWT est bien un admin.

    - Extrait l'id utilisateur depuis le token Supabase.
    - Vérifie dans public.users que is_admin = true pour cet id.
    - Retourne l'id utilisateur si admin, sinon lève HTTP 403.
    """

    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Jeton d'authentification manquant")

    token = authorization.split(" ", 1)[1].strip()
    # Log prudent: ne jamais afficher le token complet.
    # Permet de diagnostiquer si le frontend envoie bien un JWT "header.payload.signature".
    try:
        parts = token.count(".") + 1 if token else 0
        logger.warning("[auth] token received: parts=%s len=%s", parts, len(token) if token else 0)
    except Exception:
        pass
    payload = _verify_supabase_jwt(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Jeton d'authentification invalide")

    user_id = payload.get("sub") or payload.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Jeton d'authentification invalide")

    with get_db() as db:
        row = db.execute(
            text(
                """
                select is_admin
                from public.users
                where id = :user_id
                  and statut = 'actif'
                """
            ),
            {"user_id": user_id},
        ).mappings().first()

        if not row or not row.get("is_admin"):
            raise HTTPException(status_code=403, detail="Accès réservé aux administrateurs")

    return user_id


def _require_authenticated_user(authorization: str | None) -> str:
    """
    Vérifie le JWT Supabase et retourne l'id utilisateur (sans requête BDD, sans exiger is_admin).
    Utilisé pour la validation des classifications par les agents connectés.
    """
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Jeton d'authentification manquant")
    token = authorization.split(" ", 1)[1].strip()
    payload = _verify_supabase_jwt(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Jeton d'authentification invalide")
    uid = payload.get("sub") or payload.get("user_id")
    if not uid:
        raise HTTPException(status_code=401, detail="Jeton d'authentification invalide")
    return str(uid)


def user_required(authorization: str | None = Header(default=None)) -> str:
    return _require_authenticated_user(authorization)


def _create_supabase_auth_user(email: str, password: str) -> str | None:
    """
    Crée un utilisateur dans Supabase Auth via l'API d'admin.

    - Retourne l'id de l'utilisateur auth (string) en cas de succès.
    - Si la configuration Supabase Admin est absente, retourne simplement None
      sans lever d'erreur, pour conserver le comportement existant.
    """

    supabase_url = Config.SUPABASE_URL
    service_key = Config.SUPABASE_SERVICE_ROLE_KEY
    if not supabase_url or not service_key:
        # Configuration non fournie : on ne crée que dans public.users.
        return None

    url = f"{supabase_url.rstrip('/')}/auth/v1/admin/users"
    headers = {
        "apikey": service_key,
        "Authorization": f"Bearer {service_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "email": email,
        "password": password,
        "email_confirm": True,
    }

    resp = requests.post(url, headers=headers, json=payload, timeout=10)
    if not resp.ok:
        # On journalise l'erreur côté serveur mais on n'empêche pas
        # la création dans public.users, pour ne pas bloquer l'admin.
        try:
            detail = resp.json()
        except Exception:
            detail = resp.text
        logger.warning("[Supabase Auth] Echec creation user %s: %s", email, detail)
        return None

    data = resp.json()
    # id est le champ standard dans auth.users
    return data.get("id")


def _reset_supabase_auth_password(email: str, new_password: str) -> None:
    """
    Réinitialise le mot de passe d'un utilisateur Supabase Auth à partir de son email.

    Best-effort :
    - Si la configuration Supabase n'est pas fournie ou si l'appel échoue,
      on ne bloque pas la route ; on se contente de journaliser l'erreur.
    """

    supabase_url = Config.SUPABASE_URL
    service_key = Config.SUPABASE_SERVICE_ROLE_KEY
    if not supabase_url or not service_key:
        return

    base_url = supabase_url.rstrip("/")
    headers = {
        "apikey": service_key,
        "Authorization": f"Bearer {service_key}",
        "Content-Type": "application/json",
    }

    # 1. Recherche de l'utilisateur Auth par email
    try:
        list_resp = requests.get(
            f"{base_url}/auth/v1/admin/users",
            headers=headers,
            params={"email": email, "per_page": 1},
            timeout=10,
        )
    except Exception as exc:  # pragma: no cover - garde-fou réseau
        logger.exception(
            "[Supabase Auth] Erreur réseau lors de la recherche user %s",
            email,
        )
        return

    if not list_resp.ok:
        logger.warning(
            "[Supabase Auth] Echec recherche user %s: %s %s",
            email,
            list_resp.status_code,
            list_resp.text,
        )
        return

    try:
        users = list_resp.json()
    except Exception:
        logger.warning(
            "[Supabase Auth] Réponse JSON invalide lors de la recherche user %s",
            email,
        )
        return

    if isinstance(users, dict) and "users" in users:
        users_list = users.get("users") or []
    else:
        users_list = users if isinstance(users, list) else []

    if not users_list:
        logger.warning("[Supabase Auth] Aucun compte Auth trouvé pour %s", email)
        return

    auth_id = users_list[0].get("id")
    if not auth_id:
        logger.warning("[Supabase Auth] Réponse sans id pour %s", email)
        return

    # 2. Mise à jour du mot de passe pour cet id
    try:
        update_resp = requests.put(
            f"{base_url}/auth/v1/admin/users/{auth_id}",
            headers=headers,
            json={"password": new_password},
            timeout=10,
        )
    except Exception as exc:  # pragma: no cover - garde-fou réseau
        logger.exception(
            "[Supabase Auth] Erreur réseau lors de la mise à jour du mot de passe pour %s",
            email,
        )
        return

    if not update_resp.ok:
        logger.warning(
            "[Supabase Auth] Echec mise à jour mot de passe pour %s: %s %s",
            email,
            update_resp.status_code,
            update_resp.text,
        )


def _insert_audit_log(
    *,
    actor_id: str,
    action: str,
    entity_type: str,
    entity_id: str,
    details: dict | None = None,
) -> None:
    """
    Insère une entrée dans la table d'audit.

    Best-effort : en cas d'erreur, on journalise mais on ne bloque pas la requête.
    """

    from uuid import uuid4

    payload_details = details or {}
    try:
        details_json = json.dumps(payload_details, ensure_ascii=False)
    except Exception:
        # En cas de problème de sérialisation, on force un dict minimal.
        details_json = json.dumps({"error": "unserializable details"})

    now = datetime.now(timezone.utc)

    try:
        with get_db() as db:
            db.execute(
                text(
                    """
                    insert into public.audit_logs (
                      id,
                      created_at,
                      actor_id,
                      action,
                      entity_type,
                      entity_id,
                      details
                    )
                    values (
                      :id,
                      :created_at,
                      :actor_id,
                      :action,
                      :entity_type,
                      :entity_id,
                      :details
                    )
                    """
                ),
                {
                    "id": str(uuid4()),
                    "created_at": now,
                    "actor_id": actor_id,
                    "action": action,
                    "entity_type": entity_type,
                    "entity_id": entity_id,
                    "details": details_json,
                },
            )
            db.commit()
    except Exception as exc:  # pragma: no cover - garde-fou
        logger.warning(
            "[AUDIT] Echec insertion log %s sur %s:%s: %s",
            action,
            entity_type,
            entity_id,
            exc,
        )


def _delete_supabase_auth_user(email: str) -> None:
    """
    Supprime un utilisateur dans Supabase Auth à partir de son email.

    Best-effort :
    - Si la configuration Supabase n'est pas fournie ou si l'appel échoue,
      on ne bloque pas la route ; on se contente de journaliser l'erreur.
    """

    supabase_url = Config.SUPABASE_URL
    service_key = Config.SUPABASE_SERVICE_ROLE_KEY
    if not supabase_url or not service_key:
        return

    base_url = supabase_url.rstrip("/")
    headers = {
        "apikey": service_key,
        "Authorization": f"Bearer {service_key}",
        "Content-Type": "application/json",
    }

    try:
        list_resp = requests.get(
            f"{base_url}/auth/v1/admin/users",
            headers=headers,
            params={"email": email, "per_page": 1},
            timeout=10,
        )
    except Exception as exc:  # pragma: no cover - garde-fou réseau
        logger.exception(
            "[Supabase Auth] Erreur réseau lors de la recherche user %s",
            email,
        )
        return

    if not list_resp.ok:
        logger.warning(
            "[Supabase Auth] Echec recherche user %s avant suppression: %s %s",
            email,
            list_resp.status_code,
            list_resp.text,
        )
        return

    try:
        users = list_resp.json()
    except Exception:
        logger.warning(
            "[Supabase Auth] Réponse JSON invalide lors de la recherche user %s",
            email,
        )
        return

    if isinstance(users, dict) and "users" in users:
        users_list = users.get("users") or []
    else:
        users_list = users if isinstance(users, list) else []

    if not users_list:
        logger.warning(
            "[Supabase Auth] Aucun compte Auth trouvé pour %s à supprimer",
            email,
        )
        return

    auth_id = users_list[0].get("id")
    if not auth_id:
        logger.warning(
            "[Supabase Auth] Réponse sans id pour %s à la suppression",
            email,
        )
        return

    try:
        delete_resp = requests.delete(
            f"{base_url}/auth/v1/admin/users/{auth_id}",
            headers=headers,
            timeout=10,
        )
    except Exception as exc:  # pragma: no cover - garde-fou réseau
        logger.exception(
            "[Supabase Auth] Erreur réseau lors de la suppression du compte Auth pour %s",
            email,
        )
        return

    if not delete_resp.ok:
        logger.warning(
            "[Supabase Auth] Echec suppression compte Auth pour %s: %s %s",
            email,
            delete_resp.status_code,
            delete_resp.text,
        )


@app.on_event("startup")
def startup_event() -> None:
    """
    Initialise le moteur RAG (chunks + index FAISS) une seule fois au démarrage.
    Les objets sont stockés sur l'application pour être réutilisés par les endpoints.
    """

    # Assure que les tables "dossiers" existent (groupement des validations
    # par entreprise/société). Best-effort : si la DB est indisponible, on ne
    # bloque pas le démarrage.
    try:
        _ensure_dossier_tables()
    except Exception:
        logger.warning("[dossiers] impossible d'assurer les tables", exc_info=True)

    try:
        _ensure_classification_schema()
    except Exception:
        logger.warning("[schema] impossible d'assurer les colonnes classifications", exc_info=True)

    chunks, index = initialize_chatbot()
    app.state.chunks = chunks
    app.state.index = index
    tariff_label_index = build_tariff_label_index(chunks)
    set_tariff_label_index(tariff_label_index, chunks=chunks)
    app.state.tariff_label_index = tariff_label_index
    heading_narrative_index = build_heading_narrative_index(chunks)
    set_heading_narrative_index(heading_narrative_index)
    app.state.heading_narrative_index = heading_narrative_index
    tariff_rate_index = build_tariff_rate_index(chunks)
    set_tariff_rate_index(tariff_rate_index)
    app.state.tariff_rate_index = tariff_rate_index
    chapter_notes_index = build_chapter_notes_index(chunks)
    set_chapter_notes_index(chapter_notes_index)
    app.state.chapter_notes_index = chapter_notes_index
    chapter_titles_index = build_chapter_titles_index(chunks)
    set_chapter_titles_index(chapter_titles_index)
    app.state.chapter_titles_index = chapter_titles_index
    surface_sensitive_positions = build_surface_sensitive_positions(tariff_label_index)
    set_surface_sensitive_positions(surface_sensitive_positions)
    app.state.surface_sensitive_positions = surface_sensitive_positions
    logger.info("%s libelles tarifaires indexes depuis les chunks TEC", len(tariff_label_index))
    logger.info("%s grilles de taux TEC indexees depuis les chunks", len(tariff_rate_index))
    logger.info("%s chapitres avec notes TEC indexes", len(chapter_notes_index))
    logger.info("%s titres de chapitres indexes depuis les chunks TEC", len(chapter_titles_index))
    logger.info("%s positions sensibles a la surface exterieure (TEC)", len(surface_sensitive_positions))
    if product_identification_enabled():
        logger.info("Agent d'identification produit active (OpenAI)")
        if openai_web_search_enabled():
            logger.info("Recherche internet OpenAI active (Responses API web_search)")
    else:
        logger.info("Agent d'identification produit desactive")
    # Index FAISS dedie aux classifications validées (apprentissage par exemples).
    classifications_index, classifications_meta = initialize_validated_classifications_index()
    app.state.classifications_index = classifications_index
    app.state.classifications_meta = classifications_meta


def _ensure_dossier_tables() -> None:
    """
    Tables pour grouper les validations de classifications dans des "dossiers"
    (ex: Mosam Entreprise) afin que l'historique puisse afficher un dossier avec
    dedans les résultats validés.
    """
    with get_db() as db:
        db.execute(
            text(
                """
                create table if not exists public.classification_dossiers (
                    id uuid primary key default gen_random_uuid(),
                    owner_user_id uuid not null,
                    name text not null,
                    name_norm text not null,
                    description text,
                    created_at timestamptz not null default now(),
                    updated_at timestamptz not null default now()
                )
                """
            )
        )
        db.execute(
            text(
                """
                create unique index if not exists classification_dossiers_owner_name_norm_uniq
                on public.classification_dossiers (owner_user_id, name_norm)
                """
            )
        )

        # Nouveau modèle optimisé : stocker directement `dossier_id`
        # dans `public.classifications` (1 classification -> 1 dossier).
        # On ne recrée plus `classification_dossier_items`.
        try:
            db.execute(
                text(
                    """
                    alter table public.classifications
                    add column if not exists dossier_id uuid references public.classification_dossiers(id)
                    """
                )
            )
        except Exception:
            logger.warning("[dossiers] add classifications.dossier_id failed", exc_info=True)

        # Migration best-effort depuis l'ancienne table de liaison.
        try:
            exists_row = db.execute(
                text(
                    "select to_regclass('public.classification_dossier_items') is not null as exists"
                )
            ).mappings().first()
            if exists_row and exists_row.get("exists"):
                db.execute(
                    text(
                        """
                        update public.classifications c
                        set dossier_id = di.dossier_id
                        from (
                          select distinct on (classification_id)
                            classification_id,
                            dossier_id
                          from public.classification_dossier_items
                        ) di
                        where c.id = di.classification_id
                          and (c.dossier_id is null)
                        """
                    )
                )
        except Exception:
            logger.warning("[dossiers] migration items->classifications.dossier_id failed", exc_info=True)

        db.commit()


def _ensure_classification_schema() -> None:
    """Colonnes enrichies (risque, justification, identification) + correctifs schema."""
    enrichment_columns: tuple[tuple[str, str], ...] = (
        ("justification", "text"),
        ("risk_level", "text"),
        ("risk_label", "text"),
        ("position_label", "text"),
        ("classification_mode", "text"),
        ("identification_confidence", "numeric"),
        ("product_identification", "jsonb"),
        ("source_query", "text"),
    )
    with get_db() as db:
        for column_name, column_type in enrichment_columns:
            db.execute(
                text(
                    f"""
                    alter table public.classifications
                    add column if not exists {column_name} {column_type}
                    """
                )
            )
        try:
            db.execute(
                text(
                    """
                    alter table public.audit_logs
                    alter column id set default gen_random_uuid()
                    """
                )
            )
        except Exception:
            logger.debug("[schema] audit_logs id default deja present ou indisponible", exc_info=True)
        try:
            db.execute(
                text(
                    """
                    alter table public.classification_dossiers
                    add constraint classification_dossiers_owner_user_id_fkey
                    foreign key (owner_user_id) references public.users(id)
                    on delete cascade
                    """
                )
            )
        except Exception:
            logger.debug("[schema] FK classification_dossiers.owner_user_id deja presente", exc_info=True)
        db.commit()


def _ensure_dossier_id(owner_user_id: str, dossier_name: str) -> str | None:
    name = (dossier_name or "").strip()
    if not name:
        return None
    name_norm = _strip_accents_ascii(name).strip().lower()
    if not name_norm:
        return None

    with get_db() as db:
        row = db.execute(
            text(
                """
                insert into public.classification_dossiers (owner_user_id, name, name_norm)
                values (:owner_user_id, :name, :name_norm)
                on conflict (owner_user_id, name_norm)
                do update set
                    name = excluded.name,
                    updated_at = now()
                returning id::text as id
                """
            ),
            {"owner_user_id": owner_user_id, "name": name, "name_norm": name_norm},
        ).mappings().one()
        db.commit()
        return str(row.get("id")) if row else None


@app.get("/health", tags=["system"])
def health() -> dict:
    """Endpoint de healthcheck simple."""

    return {"status": "ok"}


def _clean_text_line(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s


def _filter_candidate_lines(lines: list[str], max_items: int) -> list[str]:
    """
    Filtre simple pour réduire les requêtes envoyées au LLM.
    L'objectif est d'obtenir des descriptions produit plausibles.
    """
    cleaned: list[str] = []
    stop_terms = {
        "produit",
        "marchandise",
        "qte",
        "qty",
        "quantite",
        "quantites",
        "valeur",
        "origine",
        "devise",
    }
    section_header_re = re.compile(
        r"^(?:produit|marchandise|article|designation|composition|caracteristique|"
        r"specification|usage|capacite|quantite|origine|valeur|devise)\s*:?\s*$",
        re.IGNORECASE,
    )
    for line in lines:
        line = _clean_text_line(line)
        if not line:
            continue
        if is_ui_boilerplate_line(line):
            continue
        line_norm = unicodedata.normalize("NFKD", line).encode("ascii", "ignore").decode("ascii").lower()
        line_norm = re.sub(r"\s+", " ", line_norm).strip()
        if section_header_re.match(line_norm):
            continue
        if re.fullmatch(r"\d+\s+(?:pce|pcs|pc|u|unite|kg|g|l|ml|m2|m3|eur|usd|xof|fcfa|gbp|chf)\b", line_norm):
            continue
        # Ignore les lignes purement numériques (prix/quantité isolés)
        if re.fullmatch(r"[\d\s,.\-/%]+", line_norm):
            continue
        # Ignore les libellés d'en-tête génériques isolés
        if line_norm in stop_terms:
            continue
        # Élimine les lignes trop courtes (souvent du bruit).
        if len(line) < 2:
            continue
        # Garde uniquement les lignes qui ressemblent à du contenu (lettres/ chiffres).
        if not re.search(r"[A-Za-zÀ-ÿ0-9]", line):
            continue

        # Enleve le bruit "origine/valeur/price" et decide si ce qui reste est exploitable.
        # Objectif : eviter d'envoyer des lignes de metadata seules au LLM.
        stripped = _strip_inline_metadata(line)
        if not stripped:
            continue
        if _is_noise_item_text(stripped):
            continue
        # Réduit le spam de très longues lignes.
        if len(line) > 1200:
            continue
        cleaned.append(line)
        if len(cleaned) >= max_items:
            break
    return cleaned


_PRODUCT_DOSSIER_HEADER = re.compile(
    r"^(?:produit|marchandise|article|d[eé]signation)\s*:\s*.+",
    re.IGNORECASE | re.UNICODE,
)
_QUESTION_SECTION_LINE = re.compile(r"^question\s*:?\s*$", re.IGNORECASE | re.UNICODE)
_META_QUESTION_LINE = re.compile(
    r"(?:code\s*sh|position\s*tarifaire|classif|quel\s+est\s+le\s+code)",
    re.IGNORECASE | re.UNICODE,
)
_DOSSIER_SECTION_KEYWORDS = (
    "composition",
    "caracteristique",
    "specification",
    "usage",
    "capacite",
    "quantite",
    "origine",
    "valeur",
    "devise",
)
_COMMERCIAL_METADATA_ONLY = re.compile(
    r"^(?:"
    r"(?:provenant\s+de|origine)\s+(?:la\s+|le\s+|l')?[\w\s'-]+"
    r"(?:\s+et\s+achet[eé]+e?\s+(?:a|à)\s+[\d\s.,]+\s*(?:dollars?|usd|eur|euros?|fcfa|xof)?)?"
    r"|achet[eé]+e?\s+(?:a|à)\s+[\d\s.,]+\s*(?:dollars?|usd|eur|euros?|fcfa|xof)?"
    r"|(?:valeur|prix)\s*[:=]?\s*[\d\s.,]+\s*(?:dollars?|usd|eur|euros?|fcfa|xof)?"
    r")\s*$",
    re.IGNORECASE | re.UNICODE,
)


def _is_structured_product_dossier_text(raw: str) -> bool:
    """True si le texte ressemble a une fiche Produit + Composition/Caracteristiques."""
    text = (raw or "").replace("\r", "\n").strip()
    if not text:
        return False

    first_line = ""
    for ln in text.splitlines():
        if ln.strip():
            first_line = ln.strip()
            break
    if not first_line or not _PRODUCT_DOSSIER_HEADER.match(first_line):
        return False

    ascii_norm = (
        unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii").lower()
    )
    return any(keyword in ascii_norm for keyword in _DOSSIER_SECTION_KEYWORDS)


def _is_commercial_metadata_only_text(text: str) -> bool:
    """Detecte une ligne qui ne decrit qu'une origine, une valeur ou un achat."""
    normalized = (text or "").strip()
    if not normalized:
        return False
    ascii_norm = (
        unicodedata.normalize("NFKD", normalized)
        .encode("ascii", "ignore")
        .decode("ascii")
        .lower()
    )
    ascii_norm = re.sub(r"\s+", " ", ascii_norm).strip()
    if _COMMERCIAL_METADATA_ONLY.match(ascii_norm):
        return True
    if re.search(
        r"\b(?:achet[eé]+e?\s+(?:a|à)|provenant\s+de|origine|valeur|prix|dollars?|usd|fcfa)\b",
        ascii_norm,
    ) and len(ascii_norm.split()) <= 8:
        return True
    return False


def _filter_phantom_classifications(classifications: list[Any]) -> list[dict[str, Any]]:
    """Supprime les lignes LLM qui ne decrivent pas une marchandise classifiable."""
    kept: list[dict[str, Any]] = []
    for item in classifications:
        if not isinstance(item, dict):
            continue
        description = str(item.get("description") or "").strip()
        if not description:
            continue
        if _is_commercial_metadata_only_text(description):
            continue
        hs_code = str(item.get("hs_code") or "").strip().upper()
        hs_missing = (
            not hs_code
            or "NON RENSEIGN" in hs_code
            or hs_code in ("N/A", "NA", "NON APPLICABLE")
        )
        if hs_missing and _is_commercial_metadata_only_text(description):
            continue
        if hs_missing and re.search(
            r"\b(?:achet[eé]|provenant|origine|valeur|prix|dollars?)\b",
            description,
            re.IGNORECASE,
        ) and len(description.split()) < 12:
            continue
        kept.append(item)
    return kept


def _try_parse_structured_product_dossier(raw: str) -> str | None:
    """
    Détecte une fiche produit structurée (Produit / Composition / Caractéristiques)
    et la renvoie comme une seule description à classifier.
    """
    text = (raw or "").replace("\r", "\n").strip()
    if not text:
        return None

    if not _is_structured_product_dossier_text(text):
        return None

    cleaned_lines: list[str] = []
    in_question = False
    for ln in text.splitlines():
        stripped = ln.strip()
        ln_norm = (
            unicodedata.normalize("NFKD", stripped)
            .encode("ascii", "ignore")
            .decode("ascii")
            .lower()
        )
        if _QUESTION_SECTION_LINE.match(ln_norm):
            in_question = True
            continue
        if in_question:
            if not stripped or _META_QUESTION_LINE.search(ln_norm):
                continue
            # Fin de la section question : ligne vide ou nouvelle section.
            if re.match(r"^(?:produit|composition|caract)", ln_norm):
                in_question = False
            else:
                continue
        cleaned_lines.append(ln)

    result = "\n".join(cleaned_lines).strip()
    if len(result) < 5:
        return None
    return result[:20000]


def _extract_items_from_txt(text_content: str, max_items: int) -> tuple[str, list[str]]:
    raw = (text_content or "").replace("\r", "\n").strip()
    if not raw:
        return "", []

    dossier = _try_parse_structured_product_dossier(raw)
    if dossier:
        return dossier, [dossier]

    # 0) Parsing "tableau texte" (Produit | Qté | Valeur, etc.).
    # IMPORTANT: on combine maintenant le résultat du tableau avec le reste
    # des lignes non-tabulaires (au lieu de retourner immédiatement).
    table_items, table_consumed_indexes = _extract_items_from_table_text(raw, max_items=max_items)
    table_items = table_items[:max_items]

    # 1) Priorité: découpage par blocs (souvent séparés par lignes vides)
    blocks = [b.strip() for b in re.split(r"\n\s*\n+", raw) if b.strip()]

    # 2) Si on trouve des marqueurs "Produit 1:" / "Marchandise 2:" dans un même bloc,
    #    on découpe en sous-chunks.
    marker_re = re.compile(
        r"(?im)\b(?:produit|marchandise|item)\s*(\d+)?\s*[\:\-]\s*"
    )
    chunks: list[str] = []
    for b in blocks:
        matches = list(marker_re.finditer(b))
        if len(matches) >= 2:
            for i, m in enumerate(matches):
                start = m.start()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(b)
                sub = b[start:end].strip()
                sub = marker_re.sub("", sub, count=1).strip()
                if sub:
                    chunks.append(sub)
        else:
            chunks.append(b)

    # 3) Nettoyage + réduction longueur
    items: list[str] = []
    for ch in chunks:
        ch = ch.strip()
        # Réduit les espaces sans perdre trop de contenu.
        ch = re.sub(r"\s+", " ", ch)
        if not re.search(r"[A-Za-zÀ-ÿ0-9]", ch):
            continue
        if len(ch) < 2:
            continue
        if len(ch) > 1200:
            ch = ch[:1200].strip()
        items.append(ch)
        if len(items) >= max_items:
            break

    # Fallback/complément: ligne par ligne sur les lignes NON consommées par le parsing tabulaire.
    # Ainsi, un fichier mixte (tableau + lignes libres) est traité en entier.
    remaining_lines = [
        ln for i, ln in enumerate(raw.splitlines()) if i not in table_consumed_indexes
    ]
    line_items = _filter_candidate_lines(remaining_lines, max_items=max_items)

    merged_items: list[str] = []
    for it in table_items:
        if len(merged_items) >= max_items:
            break
        merged_items.append(it)
    if len(merged_items) < max_items:
        for it in line_items:
            if len(merged_items) >= max_items:
                break
            merged_items.append(it)

    items = merged_items if merged_items else items

    if not items:
        truncated = raw[:20000]
        return truncated, [truncated] if truncated else []

    # Préfixe utile pour que split_user_queries renvoie une requête par ligne.
    effective_query = "\n".join([f"- {it}" for it in items])
    return effective_query, items


def _extract_items_from_table_text(raw_text: str, max_items: int) -> tuple[list[str], set[int]]:
    """
    Extrait des items depuis un texte à structure tabulaire.
    Exemples visés:
    - "Produit | Qté | Valeur"
    - "Ordinateur | 12 | 500000"
    """
    raw_lines = (raw_text or "").splitlines()
    if not raw_lines:
        return [], set()

    def _split_row(line: str) -> tuple[list[str], str] | None:
        for delim in ("|", ";", "\t"):
            if delim in line:
                parts = [p.strip() for p in line.split(delim)]
                if len(parts) >= 2:
                    return parts, delim
        # Fallback PDF: colonnes séparées par 2+ espaces
        parts_ws = [p.strip() for p in re.split(r"\s{2,}", line) if p.strip()]
        if len(parts_ws) >= 2:
            return parts_ws, "ws2+"
        return None

    items: list[str] = []
    consumed_line_indexes: set[int] = set()
    header_desc_idx: int | None = None
    header_qty_idx: int | None = None
    active_delim: str | None = None
    pending_desc: str | None = None

    for idx, raw_line in enumerate(raw_lines):
        line = (raw_line or "").strip()
        if not line:
            # Séparation de sections tabulaires (important pour PDF/fichiers mixtes)
            active_delim = None
            header_desc_idx = None
            header_qty_idx = None
            pending_desc = None
            continue

        parsed = _split_row(line)
        if not parsed:
            continue
        cells, delim = parsed

        cells_norm = [unicodedata.normalize("NFKD", c).encode("ascii", "ignore").decode("ascii").lower() for c in cells]
        looks_header = any(
            re.search(r"\b(produit|marchandise|description|libelle|article)\b", c)
            for c in cells_norm
        ) and any(re.search(r"\b(qte|qte\.|quantite|quantites|qty|quant)\b", c) for c in cells_norm)

        # Si on change de délimiteur mais qu'un nouvel en-tête est détecté,
        # on autorise le switch (ex: tableau '|' puis tableau ';').
        if active_delim and delim != active_delim and not looks_header:
            continue
        if not active_delim or delim != active_delim:
            active_delim = delim

        if looks_header:
            consumed_line_indexes.add(idx)
            for col_idx, c in enumerate(cells_norm):
                if header_desc_idx is None and re.search(r"\b(produit|marchandise|description|libelle|article)\b", c):
                    header_desc_idx = col_idx
                if header_qty_idx is None and re.search(r"\b(qte|qte\.|quantite|quantites|qty|quant)\b", c):
                    header_qty_idx = col_idx
            continue

        desc = ""
        qty = 1
        qty_found = False
        desc_found = False

        # Si les colonnes sont inversées ou sans en-tête, on essaie de détecter
        # la cellule "description" comme la cellule la plus textuelle.
        def _is_mostly_numeric(s: str) -> bool:
            t = (s or "").strip()
            if not t:
                return False
            return bool(re.fullmatch(r"[\d\s,.\-/%]+", t))

        if header_desc_idx is not None and header_desc_idx < len(cells):
            desc = _clean_text_line(cells[header_desc_idx])
            desc_found = bool(desc)
        else:
            non_numeric_indices = [i for i, c in enumerate(cells) if not _is_mostly_numeric(c)]
            chosen_desc_idx = non_numeric_indices[0] if non_numeric_indices else 0
            desc = _clean_text_line(cells[chosen_desc_idx])
            desc_found = bool(desc)

        if header_qty_idx is not None and header_qty_idx < len(cells):
            qmatch = re.search(r"\d+", cells[header_qty_idx])
            if qmatch:
                try:
                    qty = max(1, int(qmatch.group(0)))
                    qty_found = True
                except Exception:
                    qty = 1
        else:
            # fallback: prend la première cellule strictement numérique
            for c in cells:
                qmatch = re.search(r"^\s*\d+\s*$", c)
                if qmatch:
                    try:
                        qty = max(1, int(c.strip()))
                        qty_found = True
                        break
                    except Exception:
                        qty = 1

        # Heuristique ligne cassée (facture/proforma):
        # - ligne 1 contient description seule
        # - ligne 2 contient quantité / colonnes numériques
        if desc_found and not qty_found and len(cells) == 1:
            pending_desc = desc
            consumed_line_indexes.add(idx)
            continue

        if pending_desc and qty_found and not desc_found:
            desc = pending_desc
            desc_found = True
            pending_desc = None
            consumed_line_indexes.add(idx)

        # Pattern type "Produit: xxx Quantite: y"
        if not desc_found:
            joined = " ".join(cells)
            m_prod = re.search(r"(?i)\b(?:produit|marchandise|article)\s*[:=]\s*([^|;]+)", joined)
            if m_prod:
                desc = _clean_text_line(m_prod.group(1))
                desc_found = bool(desc)
            m_qty = re.search(r"(?i)\b(?:qte|quantite|qty)\s*[:=]?\s*(\d+)\b", joined)
            if m_qty:
                try:
                    qty = max(1, int(m_qty.group(1)))
                    qty_found = True
                except Exception:
                    qty = 1

        if desc_found and desc:
            items.append(f"{qty} {desc}" if qty > 1 else desc)
            consumed_line_indexes.add(idx)
        if len(items) >= max_items:
            break

    return items, consumed_line_indexes


def _split_leading_quantity(item: str) -> tuple[int, str, str, str, int]:
    """
    Extrait une quantité d'une ligne produit et renvoie (quantity, description).

    Cas gérés:
    - "15 ordinateurs"
    - "ordinateur de quantité 26"
    - "ordinateur quantite: 26"
    - "ordinateur qte 26"

    Si aucun motif trouvé, renvoie (1, ligne complète nettoyée, source, raw, confidence).
    """
    if not item:
        return 1, "", "explicit", "", 60
    s = item.strip()
    if not s:
        return 1, "", "explicit", "", 60

    def _invalid_quantity() -> tuple[int, str, str, str, int]:
        return 0, "", "invalid", "", 0

    # 0) Formats de lot PRIORITAIRES (avant "nombre en tête")
    # Ex: "2 cartons de 12 ordinateurs" => 24
    m_lot_head = re.match(
        r"^\s*([+-]?\d+)\s*(?:cartons?|bo[iî]tes?|caisses?|lots?|packs?|canettes?|cannettes?)\s*(?:de|x|\*)\s*([+-]?\d+)\s+(.+)$",
        s,
        re.IGNORECASE,
    )
    if m_lot_head:
        try:
            a = int(m_lot_head.group(1))
            b = int(m_lot_head.group(2))
            qty = a * b
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = (m_lot_head.group(3) or "").strip()
        return qty, text, "lot", f"{m_lot_head.group(1)}x{m_lot_head.group(2)}", 95

    m_lot_inline = re.search(
        r"(?i)\b([+-]?\d+)\s*(?:cartons?|bo[iî]tes?|caisses?|lots?|packs?|canettes?|cannettes?)\s*(?:de|x|\*)\s*([+-]?\d+)\b",
        s,
    )
    if m_lot_inline:
        try:
            a = int(m_lot_inline.group(1))
            b = int(m_lot_inline.group(2))
            qty = a * b
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = re.sub(
            r"(?i)\b[+-]?\d+\s*(?:cartons?|bo[iî]tes?|caisses?|lots?|packs?|canettes?|cannettes?)\s*(?:de|x|\*)\s*[+-]?\d+\b",
            "",
            s,
        )
        text = re.sub(r"\s+", " ", text).strip(" ,;:-")
        return qty, text, "lot", f"{m_lot_inline.group(1)}x{m_lot_inline.group(2)}", 95

    # 1) Quantité en tête: "15 ordinateurs"
    m_head = re.match(r"^\s*([+-]?\d+)\s+(.+)$", s)
    if m_head:
        try:
            qty = int(m_head.group(1))
        except Exception:
            qty = 1
        text = (m_head.group(2) or "").strip()
        if qty < 1:
            return _invalid_quantity()
        return qty, text, "explicit", m_head.group(1) or "", 95

    # 2) Formats multiplicatifs en tête:
    # "x3 ordinateur", "3x ordinateur"
    m_mult_head = re.match(r"^\s*(?:x\s*(\d+)|(\d+)\s*x)\s+(.+)$", s, re.IGNORECASE)
    if m_mult_head:
        num = m_mult_head.group(1) or m_mult_head.group(2)
        try:
            qty = int(num) if num else 1
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = (m_mult_head.group(3) or "").strip()
        return qty, text, "explicit", num or "", 95

    # 3) Quantité approximative/range en tête:
    # "~20 ordinateurs", "environ 20 ordinateurs", "20-25 ordinateurs"
    m_approx_head = re.match(
        r"^\s*(?:~|environ|approx(?:\.|imativement)?)\s*(\d+)\s+(.+)$",
        s,
        re.IGNORECASE,
    )
    if m_approx_head:
        try:
            qty = int(m_approx_head.group(1))
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = (m_approx_head.group(2) or "").strip()
        return qty, text, "explicit", m_approx_head.group(1) or "", 80

    m_range_head = re.match(r"^\s*(\d+)\s*[-/]\s*(\d+)\s+(.+)$", s)
    if m_range_head:
        try:
            a = int(m_range_head.group(1))
            b = int(m_range_head.group(2))
            # Choix prudent: borne haute du range.
            qty = max(a, b)
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = (m_range_head.group(3) or "").strip()
        return qty, text, "range_upper", f"{m_range_head.group(1)}-{m_range_head.group(2)}", 70

    # 4) Nombres en lettres (fr), ex: "vingt ordinateurs"
    number_words: dict[str, int] = {
        "zero": 0,
        "un": 1,
        "une": 1,
        "deux": 2,
        "trois": 3,
        "quatre": 4,
        "cinq": 5,
        "six": 6,
        "sept": 7,
        "huit": 8,
        "neuf": 9,
        "dix": 10,
        "onze": 11,
        "douze": 12,
        "treize": 13,
        "quatorze": 14,
        "quinze": 15,
        "seize": 16,
        "dixsept": 17,
        "dixhuit": 18,
        "dixneuf": 19,
        "vingt": 20,
        "trente": 30,
        "quarante": 40,
        "cinquante": 50,
        "soixante": 60,
        "cent": 100,
        "mille": 1000,
    }

    normalized_for_words = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    normalized_for_words = re.sub(r"[-_]", " ", normalized_for_words.lower()).strip()
    tokens = [tok for tok in normalized_for_words.split() if tok]
    compact = " ".join(tokens)
    compact = compact.replace("dix sept", "dixsept").replace("dix huit", "dixhuit").replace("dix neuf", "dixneuf")
    compact = compact.replace("quatre vingts", "quatrevingts").replace("quatre vingt", "quatrevingt")
    compact = compact.replace("soixante dix", "soixantedix").replace("soixante onze", "soixanteonze")
    compact = compact.replace("soixante douze", "soixantedouze").replace("soixante treize", "soixantetreize")
    compact = compact.replace("soixante quatorze", "soixantequatorze").replace("soixante quinze", "soixantequinze")
    compact = compact.replace("soixante seize", "soixanteseize")
    compact = compact.replace("quatre vingt dix", "quatrevingtdix")
    compact = compact.replace("quatre vingt onze", "quatrevingtonze")
    compact = compact.replace("quatre vingt douze", "quatrevingtdouze")
    compact = compact.replace("quatre vingt treize", "quatrevingttreize")
    compact = compact.replace("quatre vingt quatorze", "quatrevingtquatorze")
    compact = compact.replace("quatre vingt quinze", "quatrevingtquinze")
    compact = compact.replace("quatre vingt seize", "quatrevingtseize")
    tokens = [tok for tok in compact.split() if tok]
    special_words: dict[str, int] = {
        "soixantedix": 70,
        "soixanteonze": 71,
        "soixantedouze": 72,
        "soixantetreize": 73,
        "soixantequatorze": 74,
        "soixantequinze": 75,
        "soixanteseize": 76,
        "quatrevingt": 80,
        "quatrevingts": 80,
        "quatrevingtdix": 90,
        "quatrevingtonze": 91,
        "quatrevingtdouze": 92,
        "quatrevingttreize": 93,
        "quatrevingtquatorze": 94,
        "quatrevingtquinze": 95,
        "quatrevingtseize": 96,
    }
    if tokens:
        word_qty = 0
        consumed = 0
        # Gère patterns simples: "vingt", "cinquante", "vingt deux", "cent vingt"
        while consumed < len(tokens):
            tkn = tokens[consumed]
            if tkn in ("et", "de", "d"):
                consumed += 1
                continue
            val = special_words.get(tkn)
            if val is None:
                val = number_words.get(tkn)
            if val is None:
                break
            if val == 1000:
                if word_qty == 0:
                    word_qty = 1000
                else:
                    word_qty *= 1000
            elif val == 100:
                if word_qty == 0:
                    word_qty = 100
                else:
                    word_qty *= 100
            else:
                word_qty += val
            consumed += 1
            # On limite la lecture à 3 tokens numériques max pour éviter des faux positifs.
            if consumed >= 3:
                break
        if word_qty > 0 and consumed > 0:
            # Reconstruit la description en retirant le préfixe consommé.
            original_tokens = [tok for tok in re.split(r"\s+", s.strip()) if tok]
            desc_tokens = original_tokens[consumed:] if consumed < len(original_tokens) else []
            text = " ".join(desc_tokens).strip()
            if text:
                return word_qty, text, "word_number", " ".join(tokens[:consumed]).strip(), 78

    # 4b) Nombres en lettres en fin de ligne (quantité en suffixe).
    # Ex: "ordinateur vingt-deux unités", "ordinateur quatre-vingts", "ordinateur mille unités".
    # Principe : on essaie de parser un nombre "fr" sur les derniers tokens (hors unités optionnelles).
    unit_suffix_tokens = {
        "unite",
        "unites",
        "unites",  # doublon volontaire (tolerant)
        "piece",
        "pieces",
        "pcs",
        "exemplaire",
        "exemplaires",
        "unites",
    }
    # "tokens" et "special_words/number_words" existent encore ici car définis dans le bloc précédent.
    if tokens:
        end_idx = len(tokens)
        while end_idx > 0 and tokens[end_idx - 1] in unit_suffix_tokens:
            end_idx -= 1
        # Nécessite au moins un token produit + un token quantité.
        if end_idx >= 2:
            max_phrase_tokens = min(3, end_idx)
            # On privilégie les phrases les plus longues (ex: "vingt-deux" -> 22, pas juste "deux").
            for n in range(max_phrase_tokens, 0, -1):
                phrase_tokens = tokens[end_idx - n : end_idx]
                # Parse phrase_tokens -> word_qty
                word_qty = 0
                ok = True
                j = 0
                while j < len(phrase_tokens):
                    tkn = phrase_tokens[j]
                    if tkn in ("et", "de", "d"):
                        j += 1
                        continue
                    val = special_words.get(tkn)
                    if val is None:
                        val = number_words.get(tkn)
                    if val is None:
                        ok = False
                        break
                    if val == 1000:
                        if word_qty == 0:
                            word_qty = 1000
                        else:
                            word_qty *= 1000
                    elif val == 100:
                        if word_qty == 0:
                            word_qty = 100
                        else:
                            word_qty *= 100
                    else:
                        word_qty += val
                    j += 1
                if not ok or word_qty < 1:
                    continue
                if end_idx - n <= 0:
                    continue
                # Utilise la version ascii-normalisée pour reconstruire le texte.
                # (La normalisation finale côté `_normalize_item_key` gère ensuite les accents.)
                text = " ".join(tokens[: end_idx - n]).strip()
                if text:
                    return word_qty, text, "word_number", " ".join(phrase_tokens).strip(), 70

    # 5) Quantité exprimée dans la phrase.
    # Ex: "ordinateur de quantité 26", "quantite: 26", "qte 26", "qté 26"
    m_inline = re.search(
        r"(?i)\b(?:de\s+)?(?:quantit[eé]|qte|qt[eé])\s*[:=]?\s*(\d+)\b",
        s,
    )
    if m_inline:
        try:
            qty = int(m_inline.group(1))
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        # Retire le segment "de quantité 26" / "quantite: 26" de la description.
        text = re.sub(
            r"(?i)\b(?:de\s+)?(?:quantit[eé]|qte|qt[eé])\s*[:=]?\s*\d+\b",
            "",
            s,
        )
        text = re.sub(r"\s+", " ", text).strip(" ,;:-")
        return qty, text, "explicit", m_inline.group(1) or "", 95

    # 6) Formulations avec unités/exemplaires:
    # Ex: "ordinateur comptant 50 exemplaires", "ordinateur 50 exemplaires",
    #     "ordinateur 50 unités", "ordinateur 50 unites", "ordinateur 50 pcs"
    m_units = re.search(
        r"(?i)\b(?:comptant\s+)?(\d+)\s*(?:exemplaires?|unit[eé]s?|unites?|pcs?|pi[eè]ces?)\b",
        s,
    )
    if m_units:
        try:
            qty = int(m_units.group(1))
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = re.sub(
            r"(?i)\b(?:comptant\s+)?\d+\s*(?:exemplaires?|unit[eé]s?|unites?|pcs?|pi[eè]ces?)\b",
            "",
            s,
        )
        text = re.sub(r"\s+", " ", text).strip(" ,;:-")
        return qty, text, "explicit", m_units.group(1) or "", 92

    # 7) Multiplicatif en fin de ligne:
    # "ordinateur * 3", "ordinateur x3", "ordinateur x 3"
    # Le "x" doit être précédé d'un espace pour éviter les faux positifs
    # sur les codes produits (ex: FP2-FX20, RX500).
    m_mult_tail = re.match(r"^\s*(.+?)\s+(?:\*|x)\s*(\d+)\s*$", s, re.IGNORECASE)
    if m_mult_tail:
        try:
            qty = int(m_mult_tail.group(2))
        except Exception:
            qty = 1
        if qty < 1:
            return _invalid_quantity()
        text = (m_mult_tail.group(1) or "").strip()
        return qty, text, "explicit", m_mult_tail.group(2) or "", 95

    return 1, s, "implicit", "", 60


def _strip_inline_metadata(text: str) -> str:
    """
    Retire des méta-informations qui ne font pas partie du nom produit.
    Ex: "ordinateur, origine Chine, valeur 500000"
    """
    s = (text or "").strip()
    if not s:
        return ""
    # Enlève le conditionnement laissé par certains formats de lot
    # pour fusionner sur l'article principal.
    s = re.sub(
        r"(?i)\b(?:cartons?|bo[iî]tes?|caisses?|lots?|packs?)\s*(?:de|x|\*)?\s*[+-]?\d*\b",
        "",
        s,
    )
    s = re.sub(r"(?i)\borigine\s*[:=]?\s*[^,;|]+", "", s)
    s = re.sub(r"(?i)\b(?:valeur|prix|price)\s*[:=]?\s*[^,;|]+", "", s)
    s = re.sub(r"\s+", " ", s).strip(" ,;:-")
    return s


_COMPOSITION_BLOCK = re.compile(
    r"\b(?:composition|compos[eé](?:\s+de)?|constitu[eé](?:\s+de)?|caract[eé]ristiques?|sp[eé]cifications?)\s*:",
    re.IGNORECASE | re.UNICODE,
)

_COMMA_CONTINUATION_START = re.compile(
    r"^(?:"
    r"non\b|ni\b|sans\b|avec\b|sauf\b|contenant|comprenant|incluant|"
    r"r[eé]duction|r[eé]duit|reduit|"
    r"additionn(?:e|é)?(?:e|é)?s?\b|addition\b|"
    r"m[eê]me\b|aussi\b|"
    r"en\b|à\b|au\b|aux\b|pour\b|"
    r"type\b|mod[eè]le|format\b|qualit[eé]\b|teneur\b|"
    r"mati[eè]re\b|grasse\b|prot[eé]ines?\b|lactose\b|"
    r"extra\b|ultra\b|super\b|semi[\s-]|demi[\s-]|"
    r"nature\b|bio\b|pur\b|brut(?:e)?\b|raffin[eé]|concentr[eé]|pasteuris[eé]|"
    r"homog[eé]n[eé]is[eé]|aromatis[eé]|sucr[eé]|[eé]dulcor|ferment[eé]|acidifi[eé]|"
    r"affin[eé]|etuv[eé]|[eé]tuv[eé]|d[eé]sodoris[eé]|d[eé]graiss[eé]|d[eé]cortiqu[eé]|"
    r"torr[eé]fi[eé]|d[eé]caf[eé]in[eé]|"
    r"long\b|en\s+conserve|conserv[eé]|"
    r"de\s+qualit[eé]|d['\u2019]une\s+teneur|d['\u2019]un\s+poids|"
    r"conditionn[eé]|destin[eé]|pr[eé]sentant|"
    r"ou\b|et\b"
    r")",
    re.IGNORECASE | re.UNICODE,
)


# Premiers mots qui signalent un nouvel article après une virgule (pas un qualificatif).
_PRODUCT_HEAD_WORDS = frozenset(
    {
        "sucre",
        "lait",
        "beurre",
        "fromage",
        "creme",
        "huile",
        "farine",
        "riz",
        "ble",
        "cafe",
        "the",
        "viande",
        "poisson",
        "volaille",
        "poulet",
        "boeuf",
        "ordinateur",
        "telephone",
        "clavier",
        "ecran",
        "vehicule",
        "voiture",
        "moto",
        "acier",
        "ciment",
        "medicament",
        "parfum",
        "savon",
        "textile",
        "tissu",
        "coton",
        "vetement",
        "chaussure",
        "bijou",
        "cheval",
        "bouteille",
        "canette",
    }
)


def _segment_first_word(segment: str) -> str:
    seg_norm = (
        unicodedata.normalize("NFKD", (segment or ""))
        .encode("ascii", "ignore")
        .decode("ascii")
        .lower()
    )
    tokens = re.findall(r"[a-z0-9]+", seg_norm)
    return tokens[0] if tokens else ""


def _starts_new_product_segment(segment: str) -> bool:
    """True si le segment commence par un nom de marchandise (nouvel article)."""
    first = _segment_first_word(segment)
    return bool(first and first in _PRODUCT_HEAD_WORDS)


def _is_comma_continuation(segment: str, *, previous: str = "") -> bool:
    """
    True si le segment après une virgule ressemble à un qualificatif tarifaire
    (ex. « non concentrés », « réduction de concentré de sucre ») plutôt qu'à un nouvel article.
    """
    seg = (segment or "").strip()
    if not seg:
        return True
    if _starts_new_product_segment(seg):
        return False
    seg_norm = unicodedata.normalize("NFKD", seg).encode("ascii", "ignore").decode("ascii")
    if _COMMA_CONTINUATION_START.search(seg_norm):
        return True
    # Suite d'une description avec bloc « composition: » explicite.
    if _COMPOSITION_BLOCK.search(previous or ""):
        return True
    return False


def _split_on_commas_preserving_decimals(text: str) -> list[str]:
    """Découpe sur les virgules en préservant les décimales françaises (ex. 1,5 %)."""
    parts: list[str] = []
    current: list[str] = []
    chars = text or ""
    for i, c in enumerate(chars):
        if c == ",":
            prev_digit = i > 0 and chars[i - 1].isdigit()
            next_digit = i + 1 < len(chars) and chars[i + 1].isdigit()
            if prev_digit and next_digit:
                current.append(c)
                continue
            chunk = "".join(current).strip()
            if chunk:
                parts.append(chunk)
            current = []
        else:
            current.append(c)
    chunk = "".join(current).strip()
    if chunk:
        parts.append(chunk)
    return parts


def _split_comma_aware(text: str) -> list[str]:
    """Découpe sur les virgules en fusionnant les qualificatifs descriptifs."""
    s = (text or "").strip()
    if not s:
        return []
    # Bloc explicite « composition: … » → une seule marchandise.
    if _COMPOSITION_BLOCK.search(s):
        cleaned = _clean_text_line(s)
        return [cleaned] if cleaned else []

    parts = _split_on_commas_preserving_decimals(s)
    if len(parts) <= 1:
        cleaned = _clean_text_line(s)
        return [cleaned] if cleaned else []

    merged: list[str] = []
    current = parts[0]
    for seg in parts[1:]:
        if _is_comma_continuation(seg, previous=current):
            current = f"{current}, {seg}"
        else:
            if current:
                merged.append(_clean_text_line(current))
            current = seg
    if current:
        merged.append(_clean_text_line(current))
    return [m for m in merged if m]


def _split_multi_article_entry(text: str) -> list[str]:
    """
    Découpe une ligne contenant plusieurs articles.
    Ex: "ordinateur + telephone", "ordinateur et telephone", "ordinateur, telephone".
    Les virgules à l'intérieur d'une description (ex. « crème de lait, non concentrés,
    sans addition de sucre ») ne séparent pas les articles.
    """
    s = (text or "").strip()
    if not s:
        return []
    # Séparateurs forts : +, ;, « et » entre articles distincts.
    strong_parts = re.split(r"\s*(?:\+|;|\bet\b)\s*", s, flags=re.IGNORECASE)
    result: list[str] = []
    for part in strong_parts:
        part = part.strip()
        if not part:
            continue
        result.extend(_split_comma_aware(part))
    cleaned = [_clean_text_line(p) for p in result if _clean_text_line(p)]
    return cleaned if cleaned else [s]


def _normalize_item_key(text: str) -> str:
    """
    Normalisation légère pour agréger les variantes proches d'un même libellé.
    Exemples visés: "ordinateur" / "ordinateurs", "taxi" / "taxis".
    """
    t = (text or "").strip().lower()
    if not t:
        return ""
    # Normalise les accents + casse pour fusionner "téléphone" et "telephone".
    t = unicodedata.normalize("NFKD", t).encode("ascii", "ignore").decode("ascii")
    # Nettoyage de ponctuation simple.
    t = re.sub(r"[^\w\s-]", " ", t)
    # Répare quelques contractions mal saisies fréquentes
    # ex: "deau" -> "d eau", "duhuille" -> "d huile" (approximation prudente)
    t = re.sub(r"\bdeau\b", "d eau", t)
    t = re.sub(r"\bde([aeiouy])", r"d \1", t)
    t = re.sub(r"\s+", " ", t).strip()
    if not t:
        return ""
    if re.fullmatch(r"[\d\s,.\-/%]+", t):
        return ""

    token_aliases = _load_aliases_map()

    def _fuzzy_alias_lookup(token: str, aliases: dict[str, str]) -> str:
        """
        Approxime token -> canonical via similarite sur les clés d'alias/canon.
        Retourne `token` tel quel si aucun match assez bon.
        """
        if not token or len(token) < 4:
            return token

        # Réduction de l'espace de recherche : mêmes premières lettres + longueur proche.
        tok0 = token[0]
        candidates = [
            k
            for k in aliases.keys()
            if k
            and k[0] == tok0
            and abs(len(k) - len(token)) <= 3
            and len(k) >= 4
            and " " not in k
        ]
        if not candidates:
            return token

        best_k: str | None = None
        best_ratio = 0.0
        for k in candidates:
            ratio = SequenceMatcher(None, token, k).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_k = k

        # Seuil plus permissif pour les tokens courts (les fautes ont souvent
        # un impact proportionnel plus important sur la "ratio" de SequenceMatcher).
        local_threshold = _ALIAS_FUZZY_THRESHOLD
        if len(token) <= 8:
            # Pour les petits tokens, SequenceMatcher est plus instable.
            # Exemple: "travaux" -> "travail" (~0.714) doit matcher.
            local_threshold = 0.70

        if best_k and best_ratio >= local_threshold:
            mapped = aliases.get(best_k, token)
            # Logging prudent uniquement quand on est "au bord" du seuil,
            # pour surveiller les cas potentiellement risqués.
            if best_ratio < (local_threshold + 0.05) and mapped != token:
                logger.debug(
                    "[alias fuzzy] token=%r best_key=%r mapped=%r ratio=%.3f threshold=%.3f",
                    token,
                    best_k,
                    mapped,
                    best_ratio,
                    local_threshold,
                )
            return mapped
        return token

    def _singularize_french_word(w: str) -> str:
        if len(w) < 4:
            return w
        # Exceptions : certaines formes au pluriel ne suivent pas
        # l'heuristique "aux" -> "al".
        if w == "travaux":
            return w
        # Ex: "animaux" -> "animal" (simple heuristique)
        if w.endswith("aux") and not w.endswith("eaux"):
            return w[:-3] + "al"
        # Exclusions simples
        if w.endswith("ss") or w.endswith("us"):
            return w
        if w.endswith("s"):
            return w[:-1]
        if w.endswith("x") and not w.endswith("eaux") and not w.endswith("aux"):
            return w[:-1]
        return w

    words = t.split(" ")
    normalized_words: list[str] = []
    ignored_single_tokens = {
        "produit",
        "marchandise",
        "qte",
        "qty",
        "quantite",
        "quantites",
        "valeur",
        "origine",
    }
    for w in words:
        # Faire une singularisation avant le mapping alias/flou :
        # ex: "ordianteurs" -> "ordianteur", puis fuzzy peut matcher vers "ordinateur".
        w = _singularize_french_word(w)
        w_exact = token_aliases.get(w)
        if w_exact:
            w = w_exact
        else:
            w = _fuzzy_alias_lookup(w, token_aliases)
        normalized_words.append(w)
    normalized = " ".join(normalized_words).strip()
    # Canonisation générique sans produit en dur:
    # - normalise "A de B" -> "A B"
    # - normalise "B en A" -> "A B"
    # Objectif: fusionner des variantes de formulation, sans lister de produits.
    normalized = re.sub(
        r"\b([a-z0-9-]+(?:\s+[a-z0-9-]+){0,2})\s+d\s+([a-z0-9-]+(?:\s+[a-z0-9-]+){0,2})\b",
        r"\1 \2",
        normalized,
    )
    normalized = re.sub(
        r"\b([a-z0-9-]+(?:\s+[a-z0-9-]+){0,2})\s+en\s+([a-z0-9-]+(?:\s+[a-z0-9-]+){0,2})\b",
        r"\2 \1",
        normalized,
    )

    # Nettoyage générique des mots de liaison et déduplication de tokens adjacents.
    tokens = [tok for tok in normalized.split(" ") if tok]
    linkers = {"d", "de", "du", "des", "en", "a", "au", "aux", "la", "le", "les"}
    compact: list[str] = []
    for tok in tokens:
        if tok in linkers:
            continue
        if compact and compact[-1] == tok:
            continue
        compact.append(tok)
    normalized = " ".join(compact[:6]).strip()
    if normalized in ignored_single_tokens:
        return ""
    return normalized


def _is_noise_item_text(text: str) -> bool:
    """
    Détecte les entrées qui ne sont pas des marchandises exploitables.
    Ex: nombres isolés, pays isolés, mots génériques.
    """
    t = (text or "").strip().lower()
    if not t:
        return True
    t = unicodedata.normalize("NFKD", t).encode("ascii", "ignore").decode("ascii")
    t = re.sub(r"\s+", " ", t).strip()
    if not t:
        return True

    # Nombres/valeurs isolés
    if re.fullmatch(r"[\d\s,.\-/%]+", t):
        return True

    generic = {
        "produit",
        "marchandise",
        "qte",
        "qty",
        "quantite",
        "quantites",
        "valeur",
        "origine",
        "devise",
        "pce",
        "pcs",
        "pc",
        "u",
        "unite",
        "eur",
        "usd",
        "xof",
        "fcfa",
        "gbp",
        "chf",
    }
    if t in generic:
        return True

    if re.fullmatch(r"[a-z]{3}", t) and t in {"eur", "usd", "xof", "gbp", "chf", "cny", "jpy", "cad"}:
        return True

    if is_ui_boilerplate_line(text):
        return True

    if _is_commercial_metadata_only_text(text):
        return True

    # Pays isolés (fréquent dans les lignes "origine XXX")
    countries = {
        "chine",
        "allemagne",
        "turquie",
        "france",
        "italie",
        "japon",
        "vietnam",
        "usa",
        "etats unis",
        "cote divoire",
    }
    if t in countries:
        return True

    return False


def _extract_structured_dossier_quantity(dossier: str) -> tuple[int, str, str, int]:
    """Extrait la quantite d'une fiche Produit structuree (section « Quantite : »)."""
    qty = 1
    raw_qty = ""
    in_section = False
    for raw_line in (dossier or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.match(r"(?i)^quantit[eé]\s*:", line):
            in_section = True
            inline = line.split(":", 1)[1].strip() if ":" in line else ""
            if inline:
                match = re.match(r"(?i)^(\d+)\s*(\S+)?", inline)
                if match:
                    qty = max(1, int(match.group(1)))
                    raw_qty = inline
                in_section = False
            continue
        if in_section:
            match = re.match(r"(?i)^(\d+)\s*(\S+)?", line)
            if match:
                qty = max(1, int(match.group(1)))
                raw_qty = line
            in_section = False
    if not raw_qty and qty > 1:
        raw_qty = str(qty)
    return qty, raw_qty, "explicit", 95


def _aggregate_items_with_quantities(
    items: list[str], max_items: int
) -> tuple[list[str], dict[str, int], int, dict[str, dict[str, Any]]]:
    """
    Agrège une liste d'items en:
    - descriptions distinctes (représentation canonique)
    - quantité totale par description
    - quantité totale globale
    """
    # key normalisée -> description canonique (première rencontrée)
    key_to_display: dict[str, str] = {}
    # description canonique -> quantité totale
    item_counts: dict[str, int] = {}
    # description canonique -> méta quantité
    item_meta: dict[str, dict[str, Any]] = {}

    for raw_item in items:
        # Fiche produit structuree : conserver le dossier complet (origine, valeur, etc.).
        if _is_structured_product_dossier_text(raw_item):
            dossier_text = raw_item.strip()
            qty, raw_qty, source, qty_conf = _extract_structured_dossier_quantity(dossier_text)
            norm_key = _normalize_item_key(dossier_text)
            if not norm_key:
                continue
            display_text = dossier_text
            if norm_key not in key_to_display:
                key_to_display[norm_key] = display_text
            display_text = key_to_display[norm_key]
            item_counts[display_text] = int(item_counts.get(display_text, 0)) + int(qty)
            meta = item_meta.get(display_text) or {
                "line_count": 0,
                "explicit_count": 0,
                "implicit_count": 0,
                "range_upper_count": 0,
                "word_number_count": 0,
                "lot_count": 0,
                "quantity_raw_samples": [],
                "confidence_weighted_sum": 0.0,
                "confidence_weight_sum": 0,
            }
            meta["line_count"] += 1
            if source == "explicit":
                meta["explicit_count"] += 1
            elif source == "implicit":
                meta["implicit_count"] += 1
            elif source == "range_upper":
                meta["range_upper_count"] += 1
            elif source == "word_number":
                meta["word_number_count"] += 1
            elif source == "lot":
                meta["lot_count"] += 1
            if raw_qty and len(meta["quantity_raw_samples"]) < 3:
                meta["quantity_raw_samples"].append(raw_qty)
            meta["confidence_weighted_sum"] += float(qty_conf) * int(qty)
            meta["confidence_weight_sum"] += int(qty)
            item_meta[display_text] = meta
            continue

        sub_items = _split_multi_article_entry(raw_item)
        if not sub_items:
            continue
        for sub in sub_items:
            # 2) extraction quantité
            qty, text, source, raw_qty, qty_conf = _split_leading_quantity(sub)
            if qty < 1 or source == "invalid" or not text:
                continue
            # 3) nettoyage méta-infos (origine/valeur/etc.)
            text = _strip_inline_metadata(text)
            if not text:
                continue
            if _is_noise_item_text(text):
                continue
            # 4) normalisation de clé
            norm_key = _normalize_item_key(text)
            if not norm_key:
                continue
            if norm_key not in key_to_display:
                key_to_display[norm_key] = text.strip()
            display_text = key_to_display[norm_key]
            item_counts[display_text] = int(item_counts.get(display_text, 0)) + int(qty)
            meta = item_meta.get(display_text) or {
                "line_count": 0,
                "explicit_count": 0,
                "implicit_count": 0,
                "range_upper_count": 0,
                "word_number_count": 0,
                "lot_count": 0,
                "quantity_raw_samples": [],
                "confidence_weighted_sum": 0.0,
                "confidence_weight_sum": 0,
            }
            meta["line_count"] += 1
            if source == "explicit":
                meta["explicit_count"] += 1
            elif source == "implicit":
                meta["implicit_count"] += 1
            elif source == "range_upper":
                meta["range_upper_count"] += 1
            elif source == "word_number":
                meta["word_number_count"] += 1
            elif source == "lot":
                meta["lot_count"] += 1
            if raw_qty and len(meta["quantity_raw_samples"]) < 3:
                meta["quantity_raw_samples"].append(raw_qty)
            meta["confidence_weighted_sum"] += float(qty_conf) * int(qty)
            meta["confidence_weight_sum"] += int(qty)
            item_meta[display_text] = meta

    unique_items = list(item_counts.keys())
    if len(unique_items) > max_items:
        unique_items = unique_items[:max_items]
        item_counts = {k: item_counts[k] for k in unique_items}
        item_meta = {k: item_meta[k] for k in unique_items if k in item_meta}

    total_quantity = sum(item_counts.values())
    for label in unique_items:
        meta = item_meta.get(label, {})
        line_count = int(meta.get("line_count", 0))
        explicit_count = int(meta.get("explicit_count", 0))
        implicit_count = int(meta.get("implicit_count", 0))
        lot_count = int(meta.get("lot_count", 0))
        range_count = int(meta.get("range_upper_count", 0))
        word_count = int(meta.get("word_number_count", 0))
        source_kinds = 0
        if explicit_count > 0:
            source_kinds += 1
        if lot_count > 0:
            source_kinds += 1
        if range_count > 0:
            source_kinds += 1
        if word_count > 0:
            source_kinds += 1

        if source_kinds >= 2:
            qsource = "mixte"
        elif lot_count > 0:
            qsource = "lot"
        elif range_count > 0:
            qsource = "range_upper"
        elif word_count > 0:
            qsource = "word_number"
        elif implicit_count > 0 and explicit_count == 0:
            qsource = "implicit"
        elif line_count > 1 and explicit_count == 0:
            qsource = "repeat"
        else:
            qsource = "explicit"
        conf_sum = float(meta.get("confidence_weighted_sum", 0.0))
        conf_w = int(meta.get("confidence_weight_sum", 0))
        qconf = int(round(conf_sum / conf_w)) if conf_w > 0 else 60
        qconf = max(0, min(100, qconf))
        meta["quantity_source"] = qsource
        meta["quantity_raw"] = ", ".join(meta.get("quantity_raw_samples", []))
        meta["quantity_confidence"] = qconf
        meta["description_quality"] = assess_description_quality(
            source_text=label if _is_structured_product_dossier_text(label) else None,
            description=label,
        )
        item_meta[label] = meta

    return unique_items, item_counts, total_quantity, item_meta


def _extract_items_from_pdf(pdf_bytes: bytes, max_items: int, max_chars: int) -> tuple[str, list[str]]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    page_texts: list[str] = []
    total_chars = 0
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if not page_text:
            continue
        if total_chars + len(page_text) > max_chars:
            page_text = page_text[: max_chars - total_chars]
        page_texts.append(page_text)
        total_chars += len(page_text)
        if total_chars >= max_chars:
            break

    raw_text = "\n".join(page_texts).strip()
    if not raw_text:
        return "", []

    # Pour améliorer la cohérence entre TXT et PDF et bénéficier de la même
    # logique de segmentation (blocs, "Produit 1:", fallback ligne par ligne),
    # on réutilise directement le helper TXT sur le texte extrait du PDF.
    effective_query, items = _extract_items_from_txt(raw_text, max_items=max_items)
    return effective_query, items


_SUPPORTED_UPLOAD_EXTENSIONS = frozenset(
    {"txt", "text", "pdf", "csv", "xlsx", "xlsm", "xls", "docx", "doc"}
)

_UPLOAD_MIME_TO_EXT: dict[str, str] = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "text/csv": "csv",
    "application/csv": "csv",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
}


def _resolve_upload_extension(filename: str, content_type: str | None) -> str:
    ext = ""
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
    if ext in _SUPPORTED_UPLOAD_EXTENSIONS:
        return ext
    ct = (content_type or "").split(";", 1)[0].strip().lower()
    return _UPLOAD_MIME_TO_EXT.get(ct, ext)


def _cell_to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _normalize_header_label(label: str) -> str:
    ascii_norm = unicodedata.normalize("NFKD", label or "").encode("ascii", "ignore").decode("ascii")
    return re.sub(r"\s+", " ", ascii_norm.lower()).strip()


def _header_matches_alias(header: str, alias: str) -> bool:
    if not header or not alias:
        return False
    if header == alias:
        return True
    if len(alias) <= 3:
        return header == alias
    return alias in header


_TABULAR_FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "designation": (
        "description",
        "libelle",
        "designation",
        "produit",
        "marchandise",
        "article",
        "intitule",
        "nom",
        "details",
    ),
    "material": ("matiere", "composition", "matiere principale"),
    "usage": ("usage", "fonction", "utilisation", "emploi"),
    "characteristics": ("caracteristique", "specification", "spec"),
    "quantity": ("qte", "quantite", "qty", "quantity", "nombre"),
    "unit": ("unite", "unite de mesure", "u.s.", "us"),
    "origin": ("origine", "pays d'origine", "pays origine", "provenance", "pays"),
    "value": ("valeur", "montant", "prix"),
    "currency": ("devise", "monnaie", "currency"),
}


def _detect_tabular_columns(header_row: list[str]) -> dict[str, int | None]:
    norm_headers = [_normalize_header_label(h) for h in header_row]
    col_map: dict[str, int | None] = {field: None for field in _TABULAR_FIELD_ALIASES}
    for field, aliases in _TABULAR_FIELD_ALIASES.items():
        for i, h in enumerate(norm_headers):
            if not h:
                continue
            if any(_header_matches_alias(h, alias) for alias in aliases):
                col_map[field] = i
                break
    return col_map


def _tabular_header_detected(header_row: list[str]) -> bool:
    norm = [_normalize_header_label(c) for c in header_row]
    hits = 0
    for h in norm:
        if not h:
            continue
        for aliases in _TABULAR_FIELD_ALIASES.values():
            if any(_header_matches_alias(h, alias) for alias in aliases):
                hits += 1
                break
    has_designation = any(
        any(_header_matches_alias(norm[i] or "", alias) for alias in _TABULAR_FIELD_ALIASES["designation"])
        for i in range(len(norm))
    )
    return has_designation or hits >= 2


def _build_structured_merchandise_item(
    designation: str,
    material: str = "",
    usage: str = "",
    characteristics: str = "",
    quantity: str = "",
    unit: str = "",
    origin: str = "",
    value: str = "",
    currency: str = "",
) -> str:
    designation = _clean_text_line(designation)
    if not designation:
        return ""
    material = _clean_text_line(material)
    usage = _clean_text_line(usage)
    characteristics = _clean_text_line(characteristics)
    quantity = _clean_text_line(quantity)
    unit = _clean_text_line(unit)
    origin = _clean_text_line(origin)
    value = _clean_text_line(value)
    currency = _clean_text_line(currency)
    quantity_display = f"{quantity} {unit}".strip() if quantity or unit else ""
    if not any((material, usage, characteristics, quantity_display, origin, value, currency)):
        return designation
    lines = [f"Produit : {designation}"]
    if material:
        lines.append(f"Composition :\n- {material}")
    if usage:
        lines.append(f"Usage :\n{usage}")
    if characteristics:
        lines.append(f"Caractéristiques :\n- {characteristics}")
    if quantity_display:
        lines.append(f"Quantité :\n{quantity_display}")
    if origin:
        lines.append(f"Origine :\n{origin}")
    if value:
        value_line = f"{value} {currency}".strip() if currency else value
        lines.append(f"Valeur :\n{value_line}")
    elif currency:
        lines.append(f"Devise :\n{currency}")
    return "\n".join(lines)


def _item_text_from_tabular_row(
    row: list[str],
    col_map: dict[str, int | None] | None,
) -> str | None:
    non_empty = [c for c in row if c.strip()]
    if not non_empty:
        return None

    def cell_at(field: str) -> str:
        if not col_map:
            return ""
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return ""
        return row[idx].strip()

    if col_map:
        designation = cell_at("designation") or next((c for c in row if c.strip()), "")
        material = cell_at("material")
        usage = cell_at("usage")
        characteristics = cell_at("characteristics")
        quantity = cell_at("quantity")
        unit = cell_at("unit")
        origin = cell_at("origin")
        value = cell_at("value")
        currency = cell_at("currency")
    else:
        designation = non_empty[0]
        material = row[1].strip() if len(row) > 1 else ""
        usage = row[2].strip() if len(row) > 2 else ""
        characteristics = row[3].strip() if len(row) > 3 else ""
        quantity = row[4].strip() if len(row) > 4 else ""
        unit = row[5].strip() if len(row) > 5 else ""
        origin = row[6].strip() if len(row) > 6 else ""
        value = row[7].strip() if len(row) > 7 else ""
        currency = row[8].strip() if len(row) > 8 else ""

    item = _build_structured_merchandise_item(
        designation,
        material,
        usage,
        characteristics,
        quantity,
        unit,
        origin,
        value,
        currency,
    )
    return item or None


def _extract_items_from_tabular_rows(
    rows: list[list[Any]],
    max_items: int,
) -> tuple[str, list[str]]:
    str_rows: list[list[str]] = []
    for row in rows:
        cells = [_clean_text_line(_cell_to_str(c)) for c in row]
        if any(cells):
            str_rows.append(cells)
    if not str_rows:
        return "", []

    first = str_rows[0]
    header_detected = _tabular_header_detected(first)
    col_map = _detect_tabular_columns(first) if header_detected else None
    data_rows = str_rows[1:] if header_detected else str_rows

    items: list[str] = []
    for row in data_rows:
        item = _item_text_from_tabular_row(row, col_map)
        if item:
            items.append(item[:1200])
        if len(items) >= max_items:
            break

    effective_query = "\n".join([f"- {it}" for it in items]) if items else ""
    return effective_query, items


def _extract_items_from_xlsx(xlsx_bytes: bytes, max_items: int) -> tuple[str, list[str]]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    items: list[str] = []
    try:
        for sheet in wb.worksheets:
            rows: list[list[Any]] = []
            for row in sheet.iter_rows(values_only=True):
                rows.append(list(row))
            if not rows:
                continue
            _, sheet_items = _extract_items_from_tabular_rows(rows, max_items - len(items))
            items.extend(sheet_items)
            if len(items) >= max_items:
                break
    finally:
        wb.close()

    items = items[:max_items]
    effective_query = "\n".join([f"- {it}" for it in items]) if items else ""
    return effective_query, items


def _extract_items_from_xls(xls_bytes: bytes, max_items: int) -> tuple[str, list[str]]:
    import xlrd

    book = xlrd.open_workbook(file_contents=xls_bytes)
    items: list[str] = []
    for sheet in book.sheets():
        rows: list[list[Any]] = []
        for rx in range(sheet.nrows):
            rows.append([sheet.cell_value(rx, cx) for cx in range(sheet.ncols)])
        if not rows:
            continue
        _, sheet_items = _extract_items_from_tabular_rows(rows, max_items - len(items))
        items.extend(sheet_items)
        if len(items) >= max_items:
            break

    items = items[:max_items]
    effective_query = "\n".join([f"- {it}" for it in items]) if items else ""
    return effective_query, items


def _extract_items_from_docx(docx_bytes: bytes, max_items: int) -> tuple[str, list[str]]:
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    items: list[str] = []

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([_clean_text_line(cell.text) for cell in row.cells])
        if not rows:
            continue
        _, table_items = _extract_items_from_tabular_rows(rows, max_items - len(items))
        items.extend(table_items)
        if len(items) >= max_items:
            break

    if len(items) < max_items:
        para_text = "\n".join(p.text for p in doc.paragraphs if (p.text or "").strip())
        if para_text.strip():
            _, txt_items = _extract_items_from_txt(para_text, max_items - len(items))
            items.extend(txt_items)

    items = items[:max_items]
    effective_query = "\n".join([f"- {it}" for it in items]) if items else ""
    return effective_query, items


def _sniff_csv_dialect(sample: str) -> csv.Dialect:
    # Petit sniffing de séparateurs fréquent.
    try:
        return csv.Sniffer().sniff(sample, delimiters=[",", ";", "\t", "|"])
    except Exception:
        # Par défaut : CSV français souvent séparé par ';'.
        class _D(csv.Dialect):
            delimiter = ";"
            quotechar = '"'
            escapechar = None
            doublequote = True
            skipinitialspace = False
            lineterminator = "\n"
            quoting = csv.QUOTE_MINIMAL

        return _D()


def _extract_items_from_csv(csv_text: str, max_items: int) -> tuple[str, list[str]]:
    csv_text = (csv_text or "").strip()
    if not csv_text:
        return "", []

    sample = csv_text[:4096]
    dialect = _sniff_csv_dialect(sample)

    f = io.StringIO(csv_text)
    reader = csv.reader(f, dialect)
    rows = list(reader)
    if not rows:
        return "", []

    return _extract_items_from_tabular_rows(rows, max_items)


def _extract_items_from_json(json_text: str, max_items: int) -> tuple[str, list[str]]:
    json_text = (json_text or "").strip()
    if not json_text:
        return "", []

    try:
        obj = json.loads(json_text)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON invalide: {type(exc).__name__}") from exc

    def _extract_from_obj(item: Any) -> str | None:
        if item is None:
            return None
        if isinstance(item, str):
            return _clean_text_line(item)
        if isinstance(item, dict):
            keys_priority = [
                "description",
                "libelle",
                "libellé",
                "marchandise",
                "produit",
                "nom",
                "intitule",
                "libelle_produit",
                "libellé_produit",
                "text",
                "item",
            ]
            for k in keys_priority:
                if k in item and isinstance(item.get(k), str):
                    val = _clean_text_line(item[k])
                    return val if val else None
            # Fallback : première valeur string trouvée
            for _, v in item.items():
                if isinstance(v, str):
                    val = _clean_text_line(v)
                    return val if val else None
            return None
        # Fallback : si c'est un nombre (ou autre), on stringify
        if isinstance(item, (int, float, bool)):
            return str(item)
        return None

    items: list[str] = []
    if isinstance(obj, list):
        iterable = obj
    elif isinstance(obj, dict):
        if isinstance(obj.get("products"), list):
            iterable = obj.get("products") or []
        elif isinstance(obj.get("items"), list):
            iterable = obj.get("items") or []
        elif isinstance(obj.get("data"), list):
            iterable = obj.get("data") or []
        else:
            iterable = [obj]
    else:
        iterable = [obj]

    for it in iterable:
        val = _extract_from_obj(it)
        if val:
            items.append(val)
        if len(items) >= max_items:
            break

    effective_query = "\n".join(items)
    return effective_query, items


@app.post(
    "/classify/file",
    response_model=ClassifyFileResponse,
    tags=["classification"],
)
async def classify_file(
    file: UploadFile = File(...),
    max_items: int = 500,
    batch_size: int = 25,
    max_chars: int = 20000,
) -> ClassifyFileResponse:
    """
    Classe des produits à partir d'un fichier (txt, pdf, csv, excel, word).

    Note: le cache de classification est alimenté uniquement lors de la validation
    (voir POST /classifications/validate).
    """
    request_id = uuid.uuid4().hex[:8]

    if max_items < 1 or max_items > 500:
        raise HTTPException(status_code=400, detail="max_items doit être entre 1 et 500")
    if batch_size < 1 or batch_size > 50:
        raise HTTPException(status_code=400, detail="batch_size doit être entre 1 et 50")
    if batch_size > max_items:
        batch_size = max_items
    if max_chars < 1000 or max_chars > 200000:
        raise HTTPException(status_code=400, detail="max_chars invalide")

    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="Fichier vide")

    # Limite pour éviter les requêtes abusives / PDFs trop lourds.
    if len(raw_bytes) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Fichier trop volumineux (max 10MB)")

    filename = (file.filename or "").strip()
    ext = _resolve_upload_extension(filename, file.content_type)

    logger.debug(
        "[classify-file %s] filename=%r content_type=%r ext=%r",
        request_id,
        filename,
        file.content_type,
        ext,
    )

    # Extraction d'items à partir du fichier.
    # On garde une liste d'items "bruts", puis on va les dédupliquer
    # pour compter les quantités par produit.
    effective_query = ""
    items: list[str] = []
    try:
        if ext in {"txt", "text"}:
            text_content = raw_bytes.decode("utf-8", errors="ignore")
            effective_query, items = _extract_items_from_txt(text_content, max_items=max_items)
        elif ext == "pdf":
            effective_query, items = _extract_items_from_pdf(
                raw_bytes, max_items=max_items, max_chars=max_chars
            )
        elif ext == "csv":
            text_content = raw_bytes.decode("utf-8", errors="ignore")
            effective_query, items = _extract_items_from_csv(text_content, max_items=max_items)
        elif ext in {"xlsx", "xlsm"}:
            effective_query, items = _extract_items_from_xlsx(raw_bytes, max_items=max_items)
        elif ext == "xls":
            effective_query, items = _extract_items_from_xls(raw_bytes, max_items=max_items)
        elif ext == "docx":
            effective_query, items = _extract_items_from_docx(raw_bytes, max_items=max_items)
        elif ext == "doc":
            raise HTTPException(
                status_code=400,
                detail=(
                    "Le format Word .doc (ancien) n'est pas supporté. "
                    "Enregistrez le fichier au format .docx ou exportez en PDF/Excel."
                ),
            )
        else:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Type de fichier non supporté. Formats acceptés : "
                    "txt, pdf, csv, xlsx, xls, docx"
                ),
            )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("[classify-file %s] extraction échouée", request_id)
        raise HTTPException(status_code=500, detail=f"Extraction fichier échouée: {type(exc).__name__}") from exc

    effective_query = (effective_query or "").strip()
    if not effective_query or not items:
        raise HTTPException(
            status_code=400,
            detail="Impossible d'extraire des marchandises du fichier. Vérifie le format.",
        )

    # Comptage des quantités par description normalisée.
    unique_items, item_counts, total_quantity, item_meta = _aggregate_items_with_quantities(
        items, max_items=max_items
    )
    if not unique_items:
        raise HTTPException(
            status_code=400,
            detail="Impossible d'interpréter les marchandises du fichier (aucun item valide).",
        )
    # Re-calcul de la requête effective basée sur les items distincts.
    effective_query = "\n".join([f"- {it}" for it in unique_items])

    cache_key = _classify_cache_key(effective_query)
    cache_disabled = cache_classify_is_disabled() or len(effective_query) > 12000
    preview = effective_query.replace("\n", " ")
    preview = preview[:80] + ("…" if len(preview) > 80 else "")

    logger.debug(
        "[classify-file %s] cache_disabled=%s key=%s preview=%r items_distinct=%s quantity_total=%s",
        request_id,
        cache_disabled,
        cache_key,
        preview,
        len(unique_items),
        total_quantity,
    )

    if not cache_disabled:
        cached_raw = cache_get(cache_key)
        if cached_raw is not None:
            raw_out = _normalize_classifications_response(_ensure_json_raw(cached_raw))
            _inspect_raw_json(raw_out, request_id, "HIT")
            # Si la réponse en cache contient déjà des quantités, on les additionne,
            # sinon on considère 1 par ligne.
            classifications = _extract_classifications(raw_out)
            qty_from_raw = 0
            for cls in classifications:
                if isinstance(cls, dict):
                    q = cls.get("quantity")
                    if isinstance(q, (int, float)) and q > 0:
                        qty_from_raw += int(q)
                    else:
                        qty_from_raw += 1
                else:
                    qty_from_raw += 1
            return ClassifyFileResponse(
                raw=raw_out,
                effective_query=effective_query,
                items_count=qty_from_raw or total_quantity,
            )

    try:
        chunks = app.state.chunks
        index = app.state.index
    except AttributeError as exc:
        raise HTTPException(status_code=503, detail="Moteur RAG non initialisé") from exc

    try:
        # Traitement en batches pour pouvoir classifier un fichier contenant
        # beaucoup plus de produits qu'un seul appel LLM ne peut gérer.
        batches: list[list[str]] = [
            unique_items[i : i + batch_size] for i in range(0, len(unique_items), batch_size)
        ]

        narrative: str | None = None
        merged_classifications: list[dict[str, Any]] = []

        # Cache in-memory pour éviter de refaire le LLM plusieurs fois
        # sur des items identiques dans le même upload (important pour les gros fichiers).
        # Key = sha256(item_text), value = premier élément de `classifications`.
        single_item_cache: dict[str, dict[str, Any]] = {}

        for batch_idx, batch_items in enumerate(batches, start=1):
            batch_input = "\n".join([f"- {it}" for it in batch_items])
            logger.debug(
                "[classify-file %s] batch %s/%s size=%s",
                request_id,
                batch_idx,
                len(batches),
                len(batch_items),
            )

            batch_pipeline = _unwrap_pipeline_result(
                process_user_input(
                batch_input,
                chunks,
                index,
                validated_index=getattr(app.state, "classifications_index", None),
                validated_meta=getattr(app.state, "classifications_meta", None),
                )
            )
            normalized_batch = _finalize_classification_response(
                batch_pipeline.llm_raw,
                batch_pipeline.product_identifications,
            )
            batch_raw_out = _ensure_json_raw(normalized_batch)
            _inspect_raw_json(batch_raw_out, request_id, f"FRESH_BATCH_{batch_idx}")

            try:
                parsed_batch = json.loads(batch_raw_out)
            except Exception:
                parsed_batch = None

            if not isinstance(parsed_batch, dict):
                raise HTTPException(
                    status_code=500,
                    detail=f"Réponse LLM invalide pour le batch {batch_idx}",
                )

            if isinstance(parsed_batch.get("narrative"), str) and not narrative:
                narrative = parsed_batch["narrative"]

            batch_classes = parsed_batch.get("classifications") or []
            if isinstance(batch_classes, list):
                returned_n = len(batch_classes)
                expected_n = len(batch_items)
                logger.debug(
                    "[classify-file %s] batch %s returned %s/%s classifications",
                    request_id,
                    batch_idx,
                    returned_n,
                    expected_n,
                )

                # Si le modèle ne renvoie pas tout le batch,
                # on refait une passe produit par produit pour compléter.
                if returned_n != expected_n:
                    logger.warning(
                        "[classify-file %s] batch %s incomplete (%s/%s). Fallback per-item.",
                        request_id,
                        batch_idx,
                        returned_n,
                        expected_n,
                    )
                    merged_classifications.extend([])
                    for it_idx, it in enumerate(batch_items, start=1):
                        item_hash = hashlib.sha256(
                            (it or "").encode("utf-8", errors="ignore")
                        ).hexdigest()
                        cached_single = single_item_cache.get(item_hash)
                        if cached_single:
                            merged_classifications.append(cached_single)
                            continue

                        single_input = f"- {it}"
                        single_pipeline = _unwrap_pipeline_result(
                            process_user_input(
                            single_input,
                            chunks,
                            index,
                            validated_index=getattr(app.state, "classifications_index", None),
                            validated_meta=getattr(app.state, "classifications_meta", None),
                            )
                        )
                        normalized_single = _finalize_classification_response(
                            single_pipeline.llm_raw,
                            single_pipeline.product_identifications,
                        )
                        single_raw_out = _ensure_json_raw(normalized_single)
                        _inspect_raw_json(
                            single_raw_out,
                            request_id,
                            f"FALLBACK_BATCH_{batch_idx}_ITEM_{it_idx}",
                        )
                        try:
                            decoded_single = json.loads(single_raw_out)
                        except Exception:
                            decoded_single = None

                        if isinstance(decoded_single, dict):
                            single_classes = decoded_single.get("classifications") or []
                            if isinstance(single_classes, list) and single_classes:
                                first = single_classes[0]
                                if isinstance(first, dict):
                                    single_item_cache[item_hash] = first
                                    merged_classifications.append(first)
                                else:
                                    merged_classifications.append(
                                        {
                                            "description": it[:200],
                                            "hs_code": "Non renseigné",
                                            "section": "N/A",
                                            "section_name": "",
                                            "chapter": "N/A",
                                            "chapter_name": "",
                                            "dd_rate": "N/R",
                                            "rs_rate": "N/R",
                                            "us_unit": "",
                                            "other_taxes": "",
                                            "justification": "fallback: format classification inattendu",
                                            "excerpt": "",
                                            "origin": "Non renseigné",
                                            "value": "Non renseigné",
                                            "confidence": 0,
                                        }
                                    )
                            else:
                                # fallback vide : on pousse quand même un objet placeholder
                                merged_classifications.append(
                                    {
                                        "description": it[:200],
                                        "hs_code": "Non renseigné",
                                        "section": "N/A",
                                        "section_name": "",
                                        "chapter": "N/A",
                                        "chapter_name": "",
                                        "dd_rate": "N/R",
                                        "rs_rate": "N/R",
                                        "us_unit": "",
                                        "other_taxes": "",
                                        "justification": "fallback: modèle n'a pas renvoyé de classification",
                                        "excerpt": "",
                                        "origin": "Non renseigné",
                                        "value": "Non renseigné",
                                        "confidence": 0,
                                    }
                                )
                        else:
                            merged_classifications.append(
                                {
                                    "description": it[:200],
                                    "hs_code": "Non renseigné",
                                    "section": "N/A",
                                    "section_name": "",
                                    "chapter": "N/A",
                                    "chapter_name": "",
                                    "dd_rate": "N/R",
                                    "rs_rate": "N/R",
                                    "us_unit": "",
                                    "other_taxes": "",
                                    "justification": "fallback: réponse non parsable",
                                    "excerpt": "",
                                    "origin": "Non renseigné",
                                    "value": "Non renseigné",
                                    "confidence": 0,
                                }
                            )
                else:
                    merged_classifications.extend(batch_classes)

        # À ce stade, `merged_classifications` est aligné sur `unique_items` :
        # une classification par description distincte. On ajoute donc un champ
        # `quantity` en fonction du nombre d'occurrences dans le fichier.
        for idx, cls in enumerate(merged_classifications):
            if not isinstance(cls, dict):
                continue
            if idx >= len(unique_items):
                continue
            src_text = unique_items[idx]
            qty = int(item_counts.get(src_text, 1))
            if qty < 1:
                qty = 1
            cls.setdefault("quantity", qty)
            meta = item_meta.get(src_text, {})
            cls.setdefault("quantity_source", meta.get("quantity_source", "explicit"))
            cls.setdefault("quantity_raw", meta.get("quantity_raw", ""))
            cls.setdefault("quantity_confidence", meta.get("quantity_confidence", 60))
            cls.setdefault("description_quality", meta.get("description_quality"))
            if cls.get("description_quality") is None:
                enrich_item_description_quality(cls, source_text=src_text)
            cls["source_query"] = src_text

        # Le LLM peut parfois renvoyer des variantes quasi identiques (doublons).
        # On fusionne ces doublons pour éviter des lignes "Bouteilles d'eau" en double.
        merged_classifications = _merge_duplicate_classifications(merged_classifications)

        merged = {
            "narrative": narrative or INDICATIVE_DISCLAIMER_FR,
            "classifications": merged_classifications,
        }
        raw_out = _normalize_classifications_response(_ensure_json_raw(merged))

    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        logger.exception("[classify-file %s] batch classification failed", request_id)
        detail = f"{type(exc).__name__}: {exc}" if str(exc) else f"{type(exc).__name__}"
        raise HTTPException(status_code=500, detail=detail) from exc

    _inspect_raw_json(raw_out, request_id, "FRESH_MERGED")

    logger.debug(
        "[classify-file %s] fresh merged done raw_len=%s raw_preview=%r batches=%s items_distinct=%s quantity_total=%s",
        request_id,
        len(raw_out),
        raw_out[:80],
        (len(unique_items) + batch_size - 1) // batch_size,
        len(unique_items),
        total_quantity,
    )

    return ClassifyFileResponse(
        raw=raw_out, effective_query=effective_query, items_count=total_quantity
    )


def _structured_item_to_dossier(item: MerchandiseItem) -> str:
    """Convertit un MerchandiseItem structuré en texte de dossier pour le pipeline."""
    lines: list[str] = [f"Produit : {item.designation.strip()}"]
    if item.material.strip():
        lines.append(f"Composition :\n- {item.material.strip()}")
    if item.usage.strip():
        lines.append(f"Usage :\n{item.usage.strip()}")
    if item.characteristics.strip():
        lines.append(f"Caractéristiques :\n- {item.characteristics.strip()}")
    qty_parts = [item.quantity.strip(), item.unit.strip()]
    quantity = " ".join(part for part in qty_parts if part)
    if quantity:
        lines.append(f"Quantité :\n{quantity}")
    if item.origin.strip():
        lines.append(f"Origine :\n{item.origin.strip()}")
    if item.value.strip():
        val = item.value.strip()
        if item.currency.strip():
            val = f"{val} {item.currency.strip()}"
        lines.append(f"Valeur :\n{val}")
    elif item.currency.strip():
        lines.append(f"Devise :\n{item.currency.strip()}")
    return "\n".join(lines)


def _build_structured_inputs(
    items: list[MerchandiseItem],
) -> tuple[str, list[str], dict[str, int], dict[str, dict[str, Any]]]:
    """
    Construit les entrées de classification à partir d'items structurés.
    Retourne (classify_input, unique_items, item_counts, item_meta).
    La quantité vient directement du champ dédié — jamais parsée depuis le texte.
    """
    unique_items: list[str] = []
    item_counts: dict[str, int] = {}
    item_meta: dict[str, dict[str, Any]] = {}

    for mi in items:
        designation = mi.designation.strip()
        if not designation:
            continue

        has_detail = any([
            mi.material.strip(), mi.usage.strip(), mi.characteristics.strip(),
            mi.quantity.strip(), mi.unit.strip(),
            mi.origin.strip(), mi.value.strip(), mi.currency.strip(),
        ])
        display = _structured_item_to_dossier(mi) if has_detail else designation

        qty_str = mi.quantity.strip()
        qty = 1
        if qty_str:
            try:
                qty = max(1, int(qty_str))
            except ValueError:
                qty = 1

        unique_items.append(display)
        item_counts[display] = qty
        item_meta[display] = {
            "line_count": 1,
            "explicit_count": 1,
            "implicit_count": 0,
            "range_upper_count": 0,
            "word_number_count": 0,
            "lot_count": 0,
            "quantity_raw_samples": [qty_str] if qty_str else [],
            "confidence_weighted_sum": 95.0 * qty,
            "confidence_weight_sum": qty,
            "quantity_source": "explicit",
            "quantity_raw": qty_str,
            "quantity_confidence": 95,
        }

    if len(unique_items) == 1:
        classify_input = unique_items[0]
    elif unique_items:
        classify_input = "\n\n".join(unique_items)
    else:
        classify_input = ""
    return classify_input, unique_items, item_counts, item_meta


def _classify_text_query(
    query: str,
    *,
    request_id: str,
    structured_items: list[MerchandiseItem] | None = None,
    progress: ClassificationProgressReporter | None = None,
) -> str:
    """Exécute le pipeline de classification texte et retourne le JSON brut final."""

    if is_assistant_meta_query(query or ""):
        if progress:
            progress.start("merchandise")
            progress.complete("merchandise")
            progress.skip("identification")
            progress.skip("tec_context")
        raw_out = _normalize_classifications_response(
            build_assistant_meta_response_json(query or ""),
            progress=progress,
        )
        return _ensure_json_raw(raw_out)

    cache_key = _classify_cache_key(query)
    cache_disabled = cache_classify_is_disabled()
    preview = (query or "").strip().replace("\n", " ")
    preview = preview[:60] + ("..." if len(preview) > 60 else "")
    logger.debug(
        "[classify %s] cache_disabled=%s key=%s query_preview=%r",
        request_id,
        cache_disabled,
        cache_key,
        preview,
    )

    if not cache_disabled:
        cached_raw = cache_get(cache_key)
        if cached_raw is not None:
            if progress:
                progress.start("merchandise")
                progress.complete("merchandise")
                progress.skip("identification")
                progress.skip("tec_context")
            raw_out = _normalize_classifications_response(
                _ensure_json_raw(cached_raw),
                progress=progress,
            )
            logger.debug(
                "[classify %s] cache HIT raw_len=%s raw_preview=%r",
                request_id,
                len(raw_out),
                raw_out[:80],
            )
            _inspect_raw_json(raw_out, request_id, "HIT")
            return raw_out
        logger.debug("[classify %s] cache MISS", request_id)

    try:
        chunks = app.state.chunks
        index = app.state.index
    except AttributeError as exc:
        raise HTTPException(status_code=503, detail="Moteur RAG non initialisé") from exc

    classify_input = query
    unique_items: list[str] = []
    item_counts: dict[str, int] = {}
    item_meta: dict[str, dict[str, Any]] = {}
    if progress:
        progress.start("merchandise")

    if structured_items:
        classify_input, unique_items, item_counts, item_meta = _build_structured_inputs(
            structured_items
        )
        if not classify_input:
            classify_input = query
    else:
        try:
            _, extracted_items = _extract_items_from_txt(query, max_items=500)
            if extracted_items:
                unique_items, item_counts, _, item_meta = _aggregate_items_with_quantities(
                    extracted_items, max_items=500
                )
                if unique_items:
                    classify_input = "\n".join([f"- {it}" for it in unique_items])
        except Exception:
            classify_input = query
            unique_items = []
            item_counts = {}
            item_meta = {}

    if progress:
        progress.complete("merchandise")

    try:
        pipeline = _unwrap_pipeline_result(
            process_user_input(
                classify_input,
                chunks,
                index,
                validated_index=getattr(app.state, "classifications_index", None),
                validated_meta=getattr(app.state, "classifications_meta", None),
                progress=progress,
            )
        )
    except Exception as exc:  # pragma: no cover - garde-fou
        logger.exception("[classify %s] process_user_input failed", request_id)
        detail = f"{type(exc).__name__}: {exc}" if str(exc) else f"{type(exc).__name__}"
        raise HTTPException(status_code=500, detail=detail) from exc

    source_for_response = (
        query
        if _is_structured_product_dossier_text(query or "")
        else (classify_input or (query or ""))
    )
    result = _inject_source_query_into_llm_response(
        pipeline.llm_raw, source_for_response
    )
    result = _finalize_classification_response(
        result,
        pipeline.product_identifications,
        progress=progress,
    )
    raw_out = _ensure_json_raw(result)

    if unique_items and item_counts:
        try:
            parsed = json.loads(raw_out)
            if isinstance(parsed, dict):
                cls = parsed.get("classifications")
                if isinstance(cls, list):
                    for idx, item in enumerate(cls):
                        if not isinstance(item, dict):
                            continue
                        if idx >= len(unique_items):
                            break
                        src = unique_items[idx]
                        qty = int(item_counts.get(src, 1))
                        if qty < 1:
                            qty = 1
                        item.setdefault("quantity", qty)
                        meta = item_meta.get(src, {})
                        item.setdefault("quantity_source", meta.get("quantity_source", "explicit"))
                        item.setdefault("quantity_raw", meta.get("quantity_raw", ""))
                        item.setdefault("quantity_confidence", meta.get("quantity_confidence", 60))
                        item.setdefault("description_quality", meta.get("description_quality"))
                        if item.get("description_quality") is None:
                            enrich_item_description_quality(item, source_text=src)
                        item["source_query"] = src
                    parsed["classifications"] = _merge_duplicate_classifications(cls)
                    for item in parsed["classifications"]:
                        if isinstance(item, dict) and not item.get("source_query"):
                            item["source_query"] = classify_input or (query or "")
                    raw_out = _normalize_classifications_response(
                        _ensure_json_raw(parsed),
                        progress=progress,
                        product_identifications=pipeline.product_identifications,
                    )
        except Exception:
            pass

    logger.debug(
        "[classify %s] fresh generation done raw_len=%s raw_preview=%r",
        request_id,
        len(raw_out),
        raw_out[:80],
    )
    _inspect_raw_json(raw_out, request_id, "FRESH")
    return raw_out


def _classification_sse_response(worker) -> StreamingResponse:
    """Encapsule une classification synchrone dans un flux SSE."""

    event_queue: queue.Queue[Any] = queue.Queue()

    def emit(event: dict[str, Any]) -> None:
        event_queue.put(event)

    progress = ClassificationProgressReporter(emit)

    def run() -> None:
        try:
            result = worker(progress)
            event_queue.put({"type": "result", "payload": result})
        except HTTPException as exc:
            detail = exc.detail if isinstance(exc.detail, str) else str(exc.detail)
            event_queue.put({"type": "error", "detail": detail, "status": exc.status_code})
        except Exception as exc:  # pragma: no cover - garde-fou
            logger.exception("[classify-stream] worker failed")
            event_queue.put(
                {
                    "type": "error",
                    "detail": f"{type(exc).__name__}: {exc}" if str(exc) else type(exc).__name__,
                    "status": 500,
                }
            )
        finally:
            event_queue.put(None)

    threading.Thread(target=run, daemon=True).start()

    def generate():
        yield sse_init_event()
        while True:
            item = event_queue.get()
            if item is None:
                break
            yield sse_event(item)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _resolve_classify_query(payload: ClassifyRequest) -> str:
    """Retourne le texte query à utiliser (fourni explicitement ou reconstruit des items)."""
    if payload.query and payload.query.strip():
        return payload.query.strip()
    if payload.items:
        parts = [mi.designation.strip() for mi in payload.items if mi.designation.strip()]
        return ", ".join(parts)
    return ""


@app.post("/classify", response_model=ClassifyResponse, tags=["classification"])
def classify(payload: ClassifyRequest) -> ClassifyResponse:
    """
    Classe une ou plusieurs marchandises.

    - `items` : tableau structuré d'articles (recommandé).
    - `query` : texte libre (fallback, utilisé quand items est absent).
    """
    request_id = uuid.uuid4().hex[:8]
    query = _resolve_classify_query(payload)
    raw_out = _classify_text_query(
        query, request_id=request_id, structured_items=payload.items,
    )
    return ClassifyResponse(raw=raw_out)


@app.post("/classify/stream", tags=["classification"])
def classify_stream(payload: ClassifyRequest) -> StreamingResponse:
    """Classification avec progression SSE alignée sur le pipeline Mosam."""
    request_id = uuid.uuid4().hex[:8]
    query = _resolve_classify_query(payload)

    def worker(progress: ClassificationProgressReporter) -> dict[str, str]:
        raw_out = _classify_text_query(
            query,
            request_id=request_id,
            structured_items=payload.items,
            progress=progress,
        )
        return {"raw": raw_out}

    return _classification_sse_response(worker)


@app.post(
    "/classifications/validate",
    tags=["classification"],
)
def validate_classification(
    payload: ValidateClassificationRequest,
    request: Request,
    actor_user_id: str = Depends(user_required),
) -> dict:
    """
    Enregistre en base UNE classification choisie par un agent.

    Cette route est appelée depuis le frontend quand l'utilisateur
    clique sur "Valider cette classification" pour une ligne donnée.
    Authentification : JWT utilisateur (pas besoin d'être administrateur).
    """
    _rate_limit(request, "classification.validate")

    cache_query = payload.query
    cache_raw = payload.raw_response

    dossier_id = None
    try:
        dossier_id = (
            _ensure_dossier_id(actor_user_id, payload.dossier_name) if payload.dossier_name else None
        )
    except Exception:
        logger.warning("[dossiers] ensure dossier failed (single)", exc_info=True)

    try:
        result = _validate_classification_one(
            payload=payload,
            actor_user_id=actor_user_id,
            dossier_id=dossier_id,
            cache_query=cache_query,
            cache_raw=cache_raw,
            cache_already_set=False,
            request=request,
        )
    except OperationalError:
        logger.exception("[classifications/validate] base de données inaccessible")
        raise HTTPException(
            status_code=503,
            detail=(
                "Impossible de joindre la base de données (délai dépassé ou réseau). "
                "Vérifiez que Supabase est joignable depuis cette machine "
                "(pare-feu, VPN, port 6543 vs 5432 dans SUPABASE_DB_URL)."
            ),
        ) from None
    return result


def _validate_classification_one(
    payload: ValidateClassificationRequest,
    actor_user_id: str,
    dossier_id: str | None,
    request: Request,
    cache_query: str | None,
    cache_raw: str | None,
    cache_already_set: bool,
) -> dict:
    """
    Implémentation commune : insère une classification validée, alimente le RAG,
    et (optionnel) met à jour le cache classify (une fois par requête).
    """
    if payload.user_id is not None and str(payload.user_id) != str(actor_user_id):
        raise HTTPException(
            status_code=403,
            detail="Le champ user_id ne correspond pas au compte connecté",
        )

    now = datetime.now(timezone.utc)

    # On stocke déjà section et chapitre sous forme "numéro - libellé" côté frontend
    section_label = payload.section or "N/A"
    chapter_label = payload.chapter or "N/A"

    with get_db() as db:
        row = db.execute(
            text(
                """
                insert into public.classifications
                (description_produit,
                 section_produit,
                 chapitre_produit,
                 code_tarifaire,
                 classification_confidence,
                 quantity,
                 dd_rate,
                 rs_rate,
                 other_taxes,
                 us_unit,
                 origin,
                 value,
                 user_id,
                 statut_validation,
                 created_at,
                 dossier_id,
                 justification,
                 risk_level,
                 risk_label,
                 position_label,
                 classification_mode,
                 identification_confidence,
                 product_identification,
                 source_query)
                values (
                  :description,
                  :section,
                  :chapitre,
                  :code,
                  :confidence,
                  :quantity,
                  :dd_rate,
                  :rs_rate,
                  :other_taxes,
                  :us_unit,
                  :origin,
                  :value,
                  :user_id,
                  :statut,
                  :created_at,
                  :dossier_id,
                  :justification,
                  :risk_level,
                  :risk_label,
                  :position_label,
                  :classification_mode,
                  :identification_confidence,
                  cast(:product_identification as jsonb),
                  :source_query
                )
                returning
                  id,
                  description_produit,
                  section_produit,
                  chapitre_produit,
                  code_tarifaire,
                  classification_confidence,
                  quantity,
                  dd_rate,
                  rs_rate,
                  other_taxes,
                  us_unit,
                  origin,
                  value,
                  user_id,
                  statut_validation,
                  created_at as date_classification,
                  justification,
                  risk_level,
                  risk_label,
                  position_label,
                  classification_mode,
                  identification_confidence,
                  product_identification,
                  source_query
                """
            ),
            {
                "description": payload.description,
                "section": section_label,
                "chapitre": chapter_label,
                "code": payload.hs_code,
                "confidence": payload.confidence,
                "quantity": payload.quantity if payload.quantity is not None else 1,
                "dd_rate": payload.dd_rate,
                "rs_rate": payload.rs_rate,
                "other_taxes": payload.other_taxes,
                "us_unit": payload.us_unit,
                "origin": payload.origin,
                "value": payload.value,
                "user_id": actor_user_id,
                "statut": "validé",
                "created_at": now,
                "dossier_id": dossier_id,
                **_classification_enrichment_params(payload),
            },
        ).mappings().one()
        db.commit()

    result = dict(row)

    # Apprentissage (RAG étendu) :
    try:
        add_validated_classification_example_to_index(
            result,
            index=getattr(app.state, "classifications_index", None),
            meta=getattr(app.state, "classifications_meta", None),
        )
    except Exception:
        logger.warning(
            "[learning] impossible d'ajouter l'exemple validé à l'index RAG",
            exc_info=True,
        )

    # Mise en cache : on l'autorise une fois par requête bulk (ou une fois par call single)
    if not cache_already_set and cache_query and cache_raw:
        cache_disabled = cache_classify_is_disabled()
        if not cache_disabled:
            cache_key = _classify_cache_key(cache_query)
            raw_str = _ensure_json_raw(cache_raw)
            cache_set(cache_key, raw_str, ex=3600)

    # Audit best-effort
    user_id_for_audit = actor_user_id
    _insert_audit_log(
        actor_id=user_id_for_audit,
        action="classification.validate",
        entity_type="classification",
        entity_id=str(result.get("id", "")),
        details={
            "hs_code": payload.hs_code,
            "section": payload.section,
            "chapter": payload.chapter,
            "statut_validation": "validé",
        },
    )

    return result


@app.post(
    "/classifications/validate/bulk",
    tags=["classification"],
)
def validate_classifications_bulk(
    payload: ValidateClassificationBulkRequest,
    request: Request,
    actor_user_id: str = Depends(user_required),
) -> dict:
    """
    Validation multiple en une seule requête.
    JWT utilisateur requis (compte connecté, sans rôle admin obligatoire).
    """
    _rate_limit(request, "classification.validate.bulk", limit=200, window_seconds=60)

    if not payload.items:
        raise HTTPException(status_code=400, detail="items est requis")

    # Cache : on le fait max 1 fois par bulk.
    cache_query = payload.query
    cache_raw = payload.raw_response

    results: list[dict] = []
    errors: list[dict] = []
    cache_already_set = cache_query is not None and cache_raw is not None

    # Si l'UI n'a pas fourni query/raw_response, on accepte aussi
    # que le premier item qui les contient déclenche le cache.
    cache_set_done = False

    dossier_id_to_use: str | None = None
    if payload.dossier_name:
        try:
            dossier_id_to_use = _ensure_dossier_id(actor_user_id, payload.dossier_name)
        except Exception:
            logger.warning("[dossiers] ensure dossier failed (bulk)", exc_info=True)

    for idx, item in enumerate(payload.items):
        try:
            # Cache par item (fallback) si non fourni au niveau bulk
            q = cache_query or item.query
            raw_resp = cache_raw or item.raw_response
            already = cache_set_done
            result = _validate_classification_one(
                payload=item,
                actor_user_id=actor_user_id,
                dossier_id=dossier_id_to_use,
                cache_query=q,
                cache_raw=raw_resp,
                cache_already_set=already,
                request=request,
            )
            results.append(result)
            if (q and raw_resp) and not cache_set_done:
                cache_set_done = True
        except OperationalError:
            logger.exception("[classifications/validate/bulk] base de données inaccessible")
            raise HTTPException(
                status_code=503,
                detail=(
                    "Impossible de joindre la base de données (délai dépassé ou réseau). "
                    "Vérifiez Supabase et la variable SUPABASE_DB_URL (ports 6543 / 5432)."
                ),
            ) from None
        except HTTPException:
            raise
        except Exception as e:
            errors.append({"index": idx, "error": f"{type(e).__name__}: {e}"})
            continue

    return {
        "ok": len(errors) == 0,
        "total": len(payload.items),
        "validated": len(results),
        "errors_len": len(errors),
        "errors": errors,
        "cache_set": cache_set_done,
    }

@app.get("/admin/cache/classify/status", tags=["admin"])
def get_classify_cache_status(
    request: Request,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Retourne l'état du cache des classifications (activé / désactivé). Réservé aux admins."""
    _rate_limit(request, "admin.cache.classify.status.get")
    _ = admin_id
    return {"disabled": cache_classify_is_disabled()}


class CacheStatusUpdate(BaseModel):
    """Body pour activer/désactiver le cache des classifications."""
    disabled: bool


@app.patch("/admin/cache/classify/status", tags=["admin"])
def update_classify_cache_status(
    request: Request,
    payload: CacheStatusUpdate | None = None,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Active ou désactive le cache des classifications. Réservé aux admins."""
    _rate_limit(request, "admin.cache.classify.status.patch")
    _ = admin_id
    disabled = payload.disabled if payload else False
    if not cache_classify_set_disabled(disabled):
        raise HTTPException(
            status_code=500,
            detail="Impossible de mettre à jour l'état du cache (vérifier la configuration Redis et les droits en écriture).",
        )
    return {"disabled": disabled}


@app.delete("/admin/cache/classify", tags=["admin"])
def clear_classify_cache(
    request: Request,
    admin_id: str = Depends(admin_required),
) -> dict:
    """
    Vide le cache des réponses de classification (clés classify:*).
    Réservé aux administrateurs. Utile après mise à jour du prompt ou des documents RAG.
    """
    _rate_limit(request, "admin.cache.classify.delete")
    _ = admin_id
    deleted = cache_clear_classify()
    return {"cleared": True, "keys_deleted": deleted}


@app.get("/admin/normalization-aliases", tags=["admin"])
def list_normalization_aliases(
    request: Request,
    admin_id: str = Depends(admin_required),
) -> list[dict]:
    """Liste les alias de normalisation configurables (admin)."""
    _rate_limit(request, "admin.normalization_aliases.get")
    _ = admin_id
    _ensure_normalization_aliases_table()
    with get_db() as db:
        rows = db.execute(
            text(
                """
                select id::text as id, alias, canonical, is_active, created_at, updated_at
                from public.normalization_aliases
                order by alias asc
                """
            )
        ).mappings().all()
    return [dict(r) for r in rows]


@app.post("/admin/normalization-aliases", tags=["admin"])
def create_normalization_alias(
    payload: NormalizationAliasCreate,
    request: Request,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Crée (ou met à jour) un alias de normalisation (admin)."""
    _rate_limit(request, "admin.normalization_aliases.post")
    _ = admin_id
    alias = (payload.alias or "").strip().lower()
    canonical = (payload.canonical or "").strip().lower()
    if not alias or not canonical:
        raise HTTPException(status_code=400, detail="alias et canonical sont requis.")
    _ensure_normalization_aliases_table()
    with get_db() as db:
        row = db.execute(
            text(
                """
                insert into public.normalization_aliases (alias, canonical, is_active)
                values (:alias, :canonical, :is_active)
                on conflict (alias) do update
                  set canonical = excluded.canonical,
                      is_active = excluded.is_active,
                      updated_at = now()
                returning id::text as id, alias, canonical, is_active, created_at, updated_at
                """
            ),
            {
                "alias": alias,
                "canonical": canonical,
                "is_active": payload.is_active,
            },
        ).mappings().one()
        db.commit()
    _load_aliases_map(refresh=True)
    return dict(row)


@app.patch("/admin/normalization-aliases/{alias_id}", tags=["admin"])
def update_normalization_alias(
    alias_id: str,
    payload: NormalizationAliasUpdate,
    request: Request,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Met à jour un alias de normalisation (admin)."""
    _rate_limit(request, "admin.normalization_aliases.patch")
    _ = admin_id
    updates: list[str] = []
    params: dict[str, Any] = {"alias_id": alias_id}
    if payload.canonical is not None:
        canonical = payload.canonical.strip().lower()
        if not canonical:
            raise HTTPException(status_code=400, detail="canonical invalide.")
        updates.append("canonical = :canonical")
        params["canonical"] = canonical
    if payload.is_active is not None:
        updates.append("is_active = :is_active")
        params["is_active"] = payload.is_active
    if not updates:
        raise HTTPException(status_code=400, detail="Aucune donnée à mettre à jour.")
    _ensure_normalization_aliases_table()
    with get_db() as db:
        row = db.execute(
            text(
                f"""
                update public.normalization_aliases
                set {", ".join(updates)}, updated_at = now()
                where id::text = :alias_id
                returning id::text as id, alias, canonical, is_active, created_at, updated_at
                """
            ),
            params,
        ).mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail="Alias introuvable.")
        db.commit()
    _load_aliases_map(refresh=True)
    return dict(row)


@app.delete("/admin/normalization-aliases/{alias_id}", tags=["admin"])
def delete_normalization_alias(
    alias_id: str,
    request: Request,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Supprime un alias de normalisation (admin)."""
    _rate_limit(request, "admin.normalization_aliases.delete")
    _ = admin_id
    _ensure_normalization_aliases_table()
    with get_db() as db:
        row = db.execute(
            text(
                """
                delete from public.normalization_aliases
                where id::text = :alias_id
                returning id::text as id
                """
            ),
            {"alias_id": alias_id},
        ).mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail="Alias introuvable.")
        db.commit()
    _load_aliases_map(refresh=True)
    return {"deleted": True, "id": row.get("id")}


class DossierResponse(BaseModel):
    id: str
    name: str


@app.get("/dossiers", tags=["history"])
def list_dossiers(admin_id: str = Depends(admin_required)) -> list[DossierResponse]:
    """Liste les dossiers (entreprises) de l'admin connecté."""
    with get_db() as db:
        rows = db.execute(
            text(
                """
                select
                  id::text as id,
                  name
                from public.classification_dossiers
                where owner_user_id = :owner_user_id
                order by created_at desc
                limit 200
                """
            ),
            {"owner_user_id": admin_id},
        ).mappings().all()
        return [DossierResponse(**dict(r)) for r in rows]


@app.get("/history", tags=["history"])
def get_history(user_id: str | None = None) -> list[dict]:
    """
    Retourne l'historique des classifications depuis la base Supabase.

    Cet endpoint reste orienté "utilisateur final" (filtre par user_id éventuel)
    et ne nécessite pas de droits administrateur.
    """

    base_sql = _classification_history_select_sql()

    params: dict[str, Any] = {}
    if user_id:
        base_sql += " where c.user_id = :user_id"
        params["user_id"] = user_id

    base_sql += " order by c.created_at desc limit 1000"

    with get_db() as db:
        rows = db.execute(text(base_sql), params).mappings().all()
        return [dict(row) for row in rows]


@app.get("/history.csv", tags=["history"])
def export_history_csv(
    user_id: str | None = None,
    q: str | None = None,
    search: str | None = None,
    section: str | None = None,
    status: str | None = None,
    dossier: str | None = None,
) -> Response:
    """
    Exporte l'historique des classifications au format CSV.

    Inclut les principaux champs utilisés dans l'interface.
    """
    # Compat : certains clients envoient `q` (search global).
    search_term = (q or search or "").strip()

    base_sql = _classification_history_select_sql()

    conditions: list[str] = []
    params: dict[str, Any] = {}

    if user_id:
        conditions.append("c.user_id = :user_id")
        params["user_id"] = user_id

    if search_term:
        conditions.append(
            "(lower(c.description_produit) like :search_term or lower(c.code_tarifaire) like :search_term)"
        )
        params["search_term"] = f"%{search_term.lower()}%"

    if section and section not in {"Toutes", "Tout"}:
        conditions.append("c.section_produit = :section")
        params["section"] = section

    if status and status not in {"Tous", "Tout"}:
        conditions.append("c.statut_validation = :status")
        params["status"] = status

    if dossier and dossier not in {"Tous", "Tout"}:
        if dossier == "__none__":
            # Pas de dossier associé
            conditions.append("c.dossier_id is null")
        elif dossier == "__has__":
            # Uniquement les classifications associées à un dossier
            conditions.append("c.dossier_id is not null")
        else:
            dossier_norm = _strip_accents_ascii(dossier).strip().lower()
            if dossier_norm:
                conditions.append("d.name_norm = :dossier_norm")
                params["dossier_norm"] = dossier_norm

    if conditions:
        base_sql += " where " + " and ".join(conditions)

    base_sql += " order by c.created_at desc limit 1000"

    with get_db() as db:
        rows = db.execute(text(base_sql), params).mappings().all()
        rows = [dict(row) for row in rows]

    headers = [
        "id",
        "description_produit",
        "section_produit",
        "chapitre_produit",
        "code_tarifaire",
        "classification_confidence",
        "quantity",
        "dd_rate",
        "rs_rate",
        "other_taxes",
        "us_unit",
        "origin",
        "value",
        "statut_validation",
        "date_classification",
        *_HISTORY_CSV_EXTRA_HEADERS,
        "dossier_id",
        "dossier_name",
        "user_id",
        "agent_name",
        "agent_id",
    ]

    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(headers)
    for row in rows:
        writer.writerow([row.get(h, "") for h in headers])

    content = output.getvalue()
    output.close()

    return Response(
        content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="historique.csv"'},
    )


@app.get("/admin/history.csv", tags=["history"])
def export_admin_history_csv(
    request: Request,
    admin_id: str = Depends(admin_required),
    search: str | None = None,
    section: str | None = None,
    status: str | None = None,
    agent: str | None = None,
    dossier: str | None = None,
    date_from: str | None = None,
    date_to: str | None = None,
    limit: int = 5000,
) -> Response:
    """
    Exporte l'historique filtré (mêmes filtres que l'admin/historique).
    Réservé aux admins.
    """
    _rate_limit(request, "admin.history.csv.export")
    _ = admin_id

    base_sql = _classification_history_select_sql()

    conditions: list[str] = []
    params: dict[str, Any] = {}

    if search and search.strip():
        conditions.append(
            "("
            "lower(c.description_produit) like :search "
            "or lower(c.code_tarifaire) like :search"
            ")"
        )
        params["search"] = f"%{search.strip().lower()}%"

    if section and section not in {"Toutes", "Tout"}:
        conditions.append("c.section_produit = :section")
        params["section"] = section

    if status and status not in {"Tous", "Tout"}:
        conditions.append("c.statut_validation = :status")
        params["status"] = status

    if agent and agent not in {"Tous", "Tout"}:
        conditions.append("u.nom_user = :agent")
        params["agent"] = agent

    if dossier and dossier not in {"Tous", "Tout"}:
        if dossier == "__none__":
            conditions.append("c.dossier_id is null")
        elif dossier == "__has__":
            conditions.append("c.dossier_id is not null")
        else:
            dossier_norm = _strip_accents_ascii(dossier).strip().lower()
            if dossier_norm:
                conditions.append("d.name_norm = :dossier_norm")
                params["dossier_norm"] = dossier_norm

    if date_from and date_from.strip():
        conditions.append("c.created_at::date >= :date_from")
        params["date_from"] = date_from.strip()

    if date_to and date_to.strip():
        conditions.append("c.created_at::date <= :date_to")
        params["date_to"] = date_to.strip()

    if conditions:
        base_sql += " where " + " and ".join(conditions)

    base_sql += " order by c.created_at desc"
    params["limit"] = max(1, min(limit, 50000))
    base_sql += " limit :limit"

    rows: list[dict[str, Any]]
    with get_db() as db:
        rows = db.execute(text(base_sql), params).mappings().all()
        rows = [dict(r) for r in rows]

    headers = [
        "id",
        "description_produit",
        "section_produit",
        "chapitre_produit",
        "code_tarifaire",
        "classification_confidence",
        "quantity",
        "dd_rate",
        "rs_rate",
        "other_taxes",
        "us_unit",
        "origin",
        "value",
        "statut_validation",
        "date_classification",
        *_HISTORY_CSV_EXTRA_HEADERS,
        "dossier_id",
        "dossier_name",
        "user_id",
        "agent_name",
        "agent_id",
    ]

    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(headers)
    for row in rows:
        writer.writerow([row.get(h, "") for h in headers])

    content = output.getvalue()
    output.close()

    return Response(
        content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="historique_filtre.csv"'},
    )


@app.get("/users", tags=["users"])
def get_users(
    admin_id: str = Depends(admin_required),
    search: str | None = None,
    statut: str | None = None,
    is_admin: bool | None = None,
) -> list[dict]:
    """
    Retourne la liste des utilisateurs depuis la base Supabase.

    Filtres optionnels:
    - search : recherche texte sur nom, identifiant, email
    - statut : 'actif', 'inactif', 'supprimé'
    - is_admin : true / false
    """

    _ = admin_id

    base_sql = """
        select
          id as user_id,
          nom_user,
          identifiant_user,
          email,
          statut,
          is_admin
        from public.users
    """
    conditions: list[str] = []
    params: dict[str, Any] = {}

    if statut:
        conditions.append("statut = :statut")
        params["statut"] = statut

    if is_admin is not None:
        conditions.append("is_admin = :is_admin")
        params["is_admin"] = is_admin

    if search:
        conditions.append(
            "("
            "lower(nom_user) like :search "
            "or lower(identifiant_user) like :search "
            "or lower(email) like :search"
            ")"
        )
        params["search"] = f"%{search.strip().lower()}%"

    if conditions:
        base_sql += " where " + " and ".join(conditions)

    base_sql += " order by created_at asc"

    with get_db() as db:
        rows = db.execute(text(base_sql), params).mappings().all()
        return [dict(row) for row in rows]


@app.get("/users.csv", tags=["users"])
def export_users_csv(
    admin_id: str = Depends(admin_required),
    search: str | None = None,
    statut: str | None = None,
    is_admin: bool | None = None,
) -> Response:
    """
    Exporte la liste des utilisateurs au format CSV, avec les mêmes filtres
    que l'endpoint JSON /users.
    """

    rows = get_users(
        admin_id=admin_id,
        search=search,
        statut=statut,
        is_admin=is_admin,
    )

    headers = [
        "user_id",
        "nom_user",
        "identifiant_user",
        "email",
        "statut",
        "is_admin",
    ]

    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(headers)
    for row in rows:
        writer.writerow([row.get(h, "") for h in headers])

    content = output.getvalue()
    output.close()

    return Response(
        content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="utilisateurs.csv"'},
    )


@app.get("/audit-logs", tags=["audit"])
def get_audit_logs(
    admin_id: str = Depends(admin_required),
    actor_id: str | None = None,
    entity_type: str | None = None,
    action: str | None = None,
    q: str | None = None,
    date_from: str | None = None,
    date_to: str | None = None,
    limit: int = 200,
) -> list[dict]:
    """
    Retourne les entrées d'audit (admin uniquement), avec filtres simples.
    """

    _ = admin_id

    base_sql = """
        select
          l.id,
          l.created_at,
          l.actor_id,
          l.action,
          l.entity_type,
          l.entity_id,
          l.details::jsonb as details,
          u.nom_user as actor_name,
          coalesce(eu.nom_user, ec.description_produit) as entity_name
        from public.audit_logs l
        left join public.users u
          on u.id::text = l.actor_id
        left join public.users eu
          on l.entity_type = 'user' and eu.id::text = l.entity_id
        left join public.classifications ec
          on l.entity_type = 'classification' and ec.id::text = l.entity_id
    """
    conditions: list[str] = []
    params: dict[str, Any] = {}

    if actor_id:
        conditions.append("actor_id = :actor_id")
        params["actor_id"] = actor_id
    if entity_type:
        conditions.append("entity_type = :entity_type")
        params["entity_type"] = entity_type
    if action:
        conditions.append("action = :action")
        params["action"] = action
    if q:
        conditions.append("(actor_id ilike :q or action ilike :q or entity_id ilike :q)")
        params["q"] = f"%{q}%"
    if date_from:
        conditions.append("created_at >= :date_from")
        params["date_from"] = date_from
    if date_to:
        conditions.append("created_at <= :date_to")
        params["date_to"] = date_to

    if conditions:
        base_sql += " where " + " and ".join(conditions)

    base_sql += " order by created_at desc limit :limit"
    params["limit"] = max(1, min(limit, 500))

    with get_db() as db:
        rows = db.execute(text(base_sql), params).mappings().all()
        return [dict(row) for row in rows]


@app.get("/audit-logs.csv", tags=["audit"])
def export_audit_logs_csv(
    admin_id: str = Depends(admin_required),
    actor_id: str | None = None,
    entity_type: str | None = None,
    action: str | None = None,
    q: str | None = None,
    date_from: str | None = None,
    date_to: str | None = None,
    limit: int = 200,
) -> Response:
    """
    Exporte les entrées d'audit au format CSV, avec les mêmes filtres que /audit-logs.
    """

    rows = get_audit_logs(
        admin_id=admin_id,
        actor_id=actor_id,
        entity_type=entity_type,
        action=action,
        q=q,
        date_from=date_from,
        date_to=date_to,
        limit=limit,
    )

    headers = [
        "id",
        "created_at",
        "actor_id",
        "actor_name",
        "action",
        "entity_type",
        "entity_id",
        "entity_name",
        "details",
    ]

    import csv
    import io

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(headers)
    for row in rows:
        writer.writerow(
            [
                row.get("id", ""),
                row.get("created_at", ""),
                row.get("actor_id", ""),
                row.get("actor_name", ""),
                row.get("action", ""),
                row.get("entity_type", ""),
                row.get("entity_id", ""),
                row.get("entity_name", ""),
                row.get("details", ""),
            ]
        )

    content = output.getvalue()
    output.close()

    return Response(
        content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="audit_logs.csv"'},
    )


@app.patch("/users/{user_id}", tags=["users"])
def update_user(
    user_id: str,
    payload: UserUpdate,
    admin_id: str = Depends(admin_required),
) -> dict:
    """
    Met à jour les informations d'un utilisateur (nom, email, identifiant, rôle, statut).
    """

    # admin_id validé via `admin_required`

    # Normalisation des champs fournis
    fields_to_update: dict[str, Any] = {}
    if payload.nom_user is not None:
        nom = payload.nom_user.strip()
        if not nom:
            raise HTTPException(status_code=400, detail="Le nom complet est obligatoire.")
        fields_to_update["nom_user"] = nom
    if payload.identifiant_user is not None:
        identifiant = payload.identifiant_user.strip()
        if not identifiant:
            raise HTTPException(status_code=400, detail="L'identifiant est obligatoire.")
        if " " in identifiant:
            raise HTTPException(
                status_code=400,
                detail="L'identifiant ne doit pas contenir d'espace.",
            )
        fields_to_update["identifiant_user"] = identifiant
    if payload.email is not None:
        email = payload.email.strip().lower()
        if not email or "@" not in email:
            raise HTTPException(status_code=400, detail="L'email fourni n'est pas valide.")
        fields_to_update["email"] = email
    if payload.is_admin is not None:
        # Empêche un administrateur de se retirer lui-même ses droits.
        if user_id == admin_id and payload.is_admin is False:
            raise HTTPException(
                status_code=400,
                detail="Vous ne pouvez pas retirer vos propres droits administrateur.",
            )
        fields_to_update["is_admin"] = payload.is_admin
    if payload.statut is not None:
        # Empêche un administrateur de se désactiver ou de se supprimer lui-même.
        if user_id == admin_id and payload.statut in {"inactif", "supprimé"}:
            raise HTTPException(
                status_code=400,
                detail="Vous ne pouvez pas désactiver ou supprimer votre propre compte.",
            )
        fields_to_update["statut"] = payload.statut

    if not fields_to_update:
        raise HTTPException(status_code=400, detail="Aucune donnée à mettre à jour.")

    with get_db() as db:
        # Vérifie que l'utilisateur existe
        existing = db.execute(
            text(
                """
                select id
                from public.users
                where id = :user_id
                """
            ),
            {"user_id": user_id},
        ).first()
        if not existing:
            raise HTTPException(status_code=404, detail="Utilisateur introuvable.")

        # Vérification d'unicité email / identifiant si modifiés
        if "email" in fields_to_update or "identifiant_user" in fields_to_update:
            email = fields_to_update.get("email")
            identifiant = fields_to_update.get("identifiant_user")
            checks: list[str] = []
            params: dict[str, Any] = {"user_id": user_id}
            if email:
                checks.append("lower(email) = :email")
                params["email"] = email
            if identifiant:
                checks.append("lower(identifiant_user) = :identifiant")
                params["identifiant"] = identifiant.lower()

            if checks:
                sql = (
                    "select 1 from public.users "
                    "where (" + " or ".join(checks) + ") "
                    "and id <> :user_id limit 1"
                )
                exists = db.execute(text(sql), params).first()
                if exists:
                    raise HTTPException(
                        status_code=400,
                        detail="Cet email ou identifiant est déjà utilisé.",
                    )

        # Construction dynamique de la requête UPDATE
        set_clauses = [f"{col} = :{col}" for col in fields_to_update.keys()]
        params = dict(fields_to_update)
        params["user_id"] = user_id

        row = db.execute(
            text(
                f"""
                update public.users
                set {", ".join(set_clauses)}
                where id = :user_id
                returning
                  id as user_id,
                  nom_user,
                  identifiant_user,
                  email,
                  statut,
                  is_admin
                """
            ),
            params,
        ).mappings().one()
        db.commit()
        result = dict(row)

    # Audit best-effort : mise à jour d'utilisateur
    _insert_audit_log(
        actor_id=admin_id,
        action="user.update",
        entity_type="user",
        entity_id=str(user_id),
        details=fields_to_update,
    )

    return result


@app.delete("/users/{user_id}", tags=["users"])
def delete_user(
    user_id: str,
    admin_id: str = Depends(admin_required),
) -> dict:
    """
    Supprime complètement un utilisateur (table public.users) après avoir tenté
    de supprimer son compte dans Supabase Auth.
    """

    if user_id == admin_id:
        raise HTTPException(
            status_code=400,
            detail="Vous ne pouvez pas supprimer votre propre compte administrateur.",
        )

    with get_db() as db:
        # On récupère l'email avant suppression pour supprimer aussi dans Auth.
        existing = db.execute(
            text(
                """
                select email
                from public.users
                where id = :user_id
                """
            ),
            {"user_id": user_id},
        ).mappings().first()

        if not existing:
            raise HTTPException(status_code=404, detail="Utilisateur introuvable.")

        email = existing.get("email")

        row = db.execute(
            text(
                """
                delete from public.users
                where id = :user_id
                returning
                  id as user_id,
                  nom_user,
                  identifiant_user,
                  email
                """
            ),
            {"user_id": user_id},
        ).mappings().first()

        db.commit()

    # Suppression best-effort du compte Auth correspondant.
    if email:
        _delete_supabase_auth_user(email=email)

    # Audit best-effort : suppression d'utilisateur
    _insert_audit_log(
        actor_id=admin_id,
        action="user.delete",
        entity_type="user",
        entity_id=str(user_id),
        details={"email": email},
    )

    return dict(row) if row else {"user_id": user_id}


@app.post("/users/{user_id}/reset-password", tags=["users"])
def reset_user_password(
    user_id: str,
    admin_id: str = Depends(admin_required),
) -> ResetPasswordResponse:
    """
    Réinitialise le mot de passe d'un utilisateur.

    - Génère un nouveau mot de passe aléatoire.
    - Met à jour le compte Supabase Auth correspondant (si configuré).
    - Ne modifie pas la table public.users (les mots de passe y sont gérés par Supabase Auth).
    """

    import secrets
    import string

    # Génération d'un mot de passe robuste et lisible (6 caractères, lettres, chiffres, symboles)
    alphabet = string.ascii_letters + string.digits + "@#$%&*?!"
    new_password = "".join(secrets.choice(alphabet) for _ in range(6))

    with get_db() as db:
        row = db.execute(
            text(
                """
                select id::text as user_id, email
                from public.users
                where id = :user_id
                """
            ),
            {"user_id": user_id},
        ).mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail="Utilisateur introuvable.")

        email = row.get("email")
        if not email:
            raise HTTPException(
                status_code=400,
                detail="Impossible de réinitialiser le mot de passe : email manquant.",
            )

    # Mise à jour best-effort côté Supabase Auth
    _reset_supabase_auth_password(email=email, new_password=new_password)

    resp = ResetPasswordResponse(
        user_id=row["user_id"],
        email=email,
        new_password=new_password,
    )

    # Audit best-effort : reset mot de passe
    _insert_audit_log(
        actor_id=admin_id,
        action="user.reset_password",
        entity_type="user",
        entity_id=str(row["user_id"]),
        details={"email": email},
    )

    return resp


@app.post("/users", tags=["users"])
def create_user(
    payload: UserCreate,
    admin_id: str = Depends(admin_required),
) -> dict:
    """Crée un nouvel utilisateur dans la base Supabase."""

    # Validation / normalisation basique des champs
    nom_user = (payload.nom_user or "").strip()
    identifiant = (payload.identifiant_user or "").strip()
    email = (payload.email or "").strip().lower()

    if not nom_user:
        raise HTTPException(status_code=400, detail="Le nom complet est obligatoire.")
    if not identifiant:
        raise HTTPException(status_code=400, detail="L'identifiant est obligatoire.")
    if " " in identifiant:
        raise HTTPException(
            status_code=400,
            detail="L'identifiant ne doit pas contenir d'espace.",
        )
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="L'email fourni n'est pas valide.")

    # Vérification d'unicité email / identifiant
    with get_db() as db:
        exists = db.execute(
            text(
                """
                select 1
                from public.users
                where lower(email) = :email
                   or lower(identifiant_user) = :identifiant
                limit 1
                """
            ),
            {"email": email, "identifiant": identifiant.lower()},
        ).first()

        if exists:
            raise HTTPException(
                status_code=400,
                detail="Cet email ou identifiant est déjà utilisé.",
            )

    # On génère un mot de passe initial aléatoire (6 caractères, lettres, chiffres, symboles),
    # que l'admin pourra communiquer à l'utilisateur (qui devra ensuite le changer).
    import secrets
    import string

    alphabet = string.ascii_letters + string.digits + "@#$%&*?!"
    initial_password = "".join(secrets.choice(alphabet) for _ in range(6))

    auth_user_id: str | None = _create_supabase_auth_user(
        email=email,
        password=initial_password,
    )

    with get_db() as db:
        row = db.execute(
            text(
                """
                insert into public.users
                (nom_user, identifiant_user, email, is_admin, statut)
                values (:nom_user, :identifiant_user, :email, :is_admin, 'actif')
                returning
                  id as user_id,
                  nom_user,
                  identifiant_user,
                  email,
                  statut,
                  is_admin
                """
            ),
            {
                "nom_user": nom_user,
                "identifiant_user": identifiant,
                "email": email,
                "is_admin": payload.is_admin,
            },
        ).mappings().one()
        db.commit()
        result = dict(row)
        # On ajoute l'info d'id Auth si on l'a, à titre informatif,
        # ainsi que le mot de passe initial généré côté serveur.
        if auth_user_id is not None:
            result["auth_user_id"] = auth_user_id
        result["initial_password"] = initial_password

    # Audit best-effort : création d'utilisateur
    _insert_audit_log(
        actor_id=admin_id,
        action="user.create",
        entity_type="user",
        entity_id=str(result.get("user_id", "")),
        details={
            "nom_user": nom_user,
            "identifiant_user": identifiant,
            "email": email,
            "is_admin": payload.is_admin,
        },
    )

    return result


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
